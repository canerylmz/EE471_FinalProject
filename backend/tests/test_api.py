"""Unit tests for the TestForge Flask API."""

from datetime import date, timedelta
from io import BytesIO

from docx import Document


class FakeAIClient:
    def __init__(self, response):
        self.response = response

    def generate_json(self, prompt):
        return self.response


FORBIDDEN_TURKISH_DOCX_TERMS = [
    "Test Raporu",
    "Ölçülen",
    "Sonuç",
    "Mühendis",
    "Tarih",
    "Test Amacı",
    "Gözlemler",
]


def _assert_no_turkish_docx_terms(data):
    text = _docx_text(data)
    for term in FORBIDDEN_TURKISH_DOCX_TERMS:
        assert term not in text


def test_health(client):
    response = client.get("/api/health")
    assert response.status_code == 200
    body = response.get_json()
    assert body["success"] is True
    assert body["data"]["status"] == "ok"


def test_create_dut_success(client):
    payload = {
        "name": "Body Control Module",
        "manufacturer": "Acme Automotive",
        "part_number": "BCM-001",
        "mounting_location": "Passenger compartment",
        "power_class": "Class II",
        "nominal_voltage": "12V",
        "customer": "ACME Corp",
        "project": "Project X",
        "temp_min": -40,
        "temp_max": 85,
        "ip_class": "IP54",
        "notes": "",
    }
    response = client.post("/api/dut", json=payload)
    assert response.status_code == 201
    body = response.get_json()
    assert body["success"] is True
    assert body["data"]["name"] == "Body Control Module"
    assert body["data"]["id"] is not None


def test_create_dut_missing_name(client):
    response = client.post("/api/dut", json={"manufacturer": "Acme"})
    assert response.status_code == 400
    body = response.get_json()
    assert body["success"] is False
    assert "error" in body


def test_list_and_get_dut(client, sample_dut):
    list_response = client.get("/api/dut")
    assert list_response.status_code == 200
    duts = list_response.get_json()["data"]
    assert len(duts) == 1
    assert duts[0]["name"] == sample_dut["name"]

    get_response = client.get(f"/api/dut/{sample_dut['id']}")
    assert get_response.status_code == 200
    assert get_response.get_json()["data"]["id"] == sample_dut["id"]


def test_delete_dut_removes_related_workflow_records(app, client, sample_dut):
    plan_response = client.post("/api/plan/generate", json={"dut_id": sample_dut["id"]})
    test_id = plan_response.get_json()["data"]["tests"][0]["id"]
    client.post(
        "/api/result",
        json={
            "test_id": test_id,
            "dut_id": sample_dut["id"],
            "measured_values": {
                "us_min": 9,
                "us_max": 16,
                "duration": 2,
                "functional_observation": "Normal operation observed.",
            },
        },
    )
    client.post(
        "/api/attachments/upload",
        data={
            "dut_id": str(sample_dut["id"]),
            "test_id": str(test_id),
            "attachment_type": "measurement_file",
            "file": (BytesIO(b"demo"), "delete_check.csv"),
        },
        content_type="multipart/form-data",
    )

    response = client.delete(f"/api/dut/{sample_dut['id']}")

    assert response.status_code == 200
    assert client.get(f"/api/dut/{sample_dut['id']}").status_code == 404
    assert app.db.query_one("SELECT COUNT(*) AS count FROM tests WHERE dut_id = ?", (sample_dut["id"],))["count"] == 0
    assert app.db.query_one("SELECT COUNT(*) AS count FROM results WHERE dut_id = ?", (sample_dut["id"],))["count"] == 0
    assert app.db.query_one("SELECT COUNT(*) AS count FROM test_plan_items WHERE dut_id = ?", (sample_dut["id"],))["count"] == 0
    assert app.db.query_one("SELECT COUNT(*) AS count FROM test_attachments WHERE dut_id = ?", (sample_dut["id"],))["count"] == 0


def test_get_dut_not_found(client):
    response = client.get("/api/dut/9999")
    assert response.status_code == 404
    assert response.get_json()["success"] is False


def test_generate_plan_returns_tests(client, sample_dut):
    response = client.post("/api/plan/generate", json={"dut_id": sample_dut["id"]})
    assert response.status_code == 200
    body = response.get_json()
    assert body["success"] is True
    tests = body["data"]["tests"]
    assert isinstance(tests, list)
    assert len(tests) > 0

    first_test = tests[0]
    for field in (
        "test_name",
        "standard_reference",
        "category",
        "status",
        "duration_hours",
        "required_equipment",
        "acceptance_criteria",
        "severity_level",
    ):
        assert field in first_test


def test_generate_plan_matches_iso_catalog(app, client, sample_dut):
    app.plan_service.ai_client = FakeAIClient(
        [
            {
                "test_name": "Overvoltage",
                "iso_part": "ISO16750-2",
                "clause_no": "4.3",
                "category": "electrical",
                "operating_mode": "Powered/operating",
                "functional_status": "Class C or project-defined",
                "required_test_level": "TBD by nominal voltage class",
                "severity": "I",
                "sample_size": "TBD",
                "reason": "12V DUT requires overvoltage verification.",
            }
        ]
    )

    response = client.post("/api/plan/generate", json={"dut_id": sample_dut["id"]})

    assert response.status_code == 200
    body = response.get_json()["data"]
    assert body["source"] == "ai-backend"
    test = body["tests"][0]
    assert test["iso_catalog_id"] is not None
    assert test["iso_part"] == "ISO16750-2"
    assert test["clause_no"] == "4.3"
    assert test["selection_reason"] == "12V DUT requires overvoltage verification."
    assert test["test_plan_item_id"] is not None
    assert test["plan_item_status"] == "planned"

    plan_items = app.db.query(
        "SELECT * FROM test_plan_items WHERE dut_id = ?", (sample_dut["id"],)
    )
    assert len(plan_items) == 1
    assert plan_items[0]["test_id"] == test["id"]


def test_generate_plan_invalid_ai_response_uses_catalog_fallback(app, client, sample_dut):
    app.plan_service.ai_client = FakeAIClient({"invalid": "shape"})

    response = client.post("/api/plan/generate", json={"dut_id": sample_dut["id"]})

    assert response.status_code == 200
    body = response.get_json()["data"]
    assert body["source"] == "fallback"
    assert body["tests"]
    assert any(test["iso_catalog_id"] is not None for test in body["tests"])

    logs = app.db.query("SELECT * FROM ai_decision_logs WHERE dut_id = ?", (sample_dut["id"],))
    assert logs
    assert logs[-1]["fallback_used"] == 1


def test_plan_item_status_update(client, sample_dut):
    plan_response = client.post("/api/plan/generate", json={"dut_id": sample_dut["id"]})
    test = plan_response.get_json()["data"]["tests"][0]

    response = client.patch(
        f"/api/plan/item/{test['test_plan_item_id']}/status",
        json={"status": "approved"},
    )

    assert response.status_code == 200
    item = response.get_json()["data"]
    assert item["status"] == "approved"

    plan_response = client.get(f"/api/plan/{sample_dut['id']}")
    updated_test = plan_response.get_json()["data"]["tests"][0]
    assert updated_test["plan_item_status"] == "approved"


def test_plan_item_order_update(client, sample_dut):
    plan_response = client.post("/api/plan/generate", json={"dut_id": sample_dut["id"]})
    test = plan_response.get_json()["data"]["tests"][0]

    response = client.patch(
        f"/api/plan/item/{test['test_plan_item_id']}/order",
        json={"sort_order": 99},
    )

    assert response.status_code == 200
    assert response.get_json()["data"]["sort_order"] == 99


def test_legacy_test_without_plan_item_is_listed(app, client, sample_dut):
    test_id = app.db.insert(
        """
        INSERT INTO tests (
            dut_id, test_name, standard_reference, category, status,
            duration_hours, required_equipment, acceptance_criteria, severity_level
        ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)
        """,
        (
            sample_dut["id"],
            "Legacy Test",
            "ISO16750 legacy",
            "electrical",
            "Mandatory",
            1,
            "[]",
            "Legacy criteria",
            "I",
        ),
    )

    response = client.get(f"/api/plan/{sample_dut['id']}")

    assert response.status_code == 200
    tests = response.get_json()["data"]["tests"]
    legacy = next(test for test in tests if test["id"] == test_id)
    assert legacy["test_plan_item_id"] is None
    assert legacy["plan_item_status"] is None


def test_generate_plan_with_empty_catalog_keeps_legacy_fallback(app, client, sample_dut):
    app.plan_service.ai_client = FakeAIClient({"invalid": "shape"})
    app.db.execute("DELETE FROM iso_test_catalog")

    response = client.post("/api/plan/generate", json={"dut_id": sample_dut["id"]})

    assert response.status_code == 200
    body = response.get_json()["data"]
    assert body["source"] == "fallback"
    assert body["tests"]
    assert all(test["iso_catalog_id"] is None for test in body["tests"])


def test_generate_plan_missing_dut(client):
    response = client.post("/api/plan/generate", json={"dut_id": 9999})
    assert response.status_code == 404
    assert response.get_json()["success"] is False


def test_checklist_generate(client, sample_dut):
    plan_response = client.post("/api/plan/generate", json={"dut_id": sample_dut["id"]})
    test_id = plan_response.get_json()["data"]["tests"][0]["id"]

    response = client.post(
        "/api/checklist/generate",
        json={"dut_id": sample_dut["id"], "test_id": test_id},
    )
    assert response.status_code == 200
    checklist = response.get_json()["data"]["checklist"]
    for key in ("equipment_calibration", "safety_precautions", "dut_preparation"):
        assert key in checklist
        assert isinstance(checklist[key], list)
        assert len(checklist[key]) > 0


def test_save_and_get_result(client, sample_dut):
    plan_response = client.post("/api/plan/generate", json={"dut_id": sample_dut["id"]})
    test_id = plan_response.get_json()["data"]["tests"][0]["id"]

    payload = {
        "test_id": test_id,
        "dut_id": sample_dut["id"],
        "result": "Pass",
        "measured_values": {"Voltage Drop": "0.2V"},
        "test_conditions": {"temperature": 23, "voltage": 12, "humidity": 45},
        "observations": "No anomalies observed.",
        "has_deviation": False,
        "engineer_name": "Jane Engineer",
    }
    save_response = client.post("/api/result", json=payload)
    assert save_response.status_code == 201
    saved = save_response.get_json()["data"]
    assert saved["result"] == "Pass"
    assert saved["measured_values"] == {"Voltage Drop": "0.2V"}

    get_response = client.get(f"/api/result/{test_id}")
    assert get_response.status_code == 200
    assert get_response.get_json()["data"]["result"] == "Pass"


def test_result_schema_endpoint(client, sample_dut):
    plan_response = client.post("/api/plan/generate", json={"dut_id": sample_dut["id"]})
    test_id = plan_response.get_json()["data"]["tests"][0]["id"]

    response = client.get(f"/api/result/schema/{test_id}")

    assert response.status_code == 200
    data = response.get_json()["data"]
    assert data["test"]["test_name"]
    assert isinstance(data["schema"], list)
    assert isinstance(data["evaluation_schema"], dict)
    assert any(field["name"] == "us_min" for field in data["schema"])


def test_dynamic_measured_values_are_saved(client, sample_dut):
    plan_response = client.post("/api/plan/generate", json={"dut_id": sample_dut["id"]})
    test_id = plan_response.get_json()["data"]["tests"][0]["id"]

    payload = {
        "test_id": test_id,
        "dut_id": sample_dut["id"],
        "result": "Pass",
        "measured_values": {
            "us_min": 9,
            "us_max": 16,
            "duration": 2,
            "functional_observation": "No reset observed.",
        },
    }
    response = client.post("/api/result", json=payload)

    assert response.status_code == 201
    saved = response.get_json()["data"]
    assert saved["measured_values"]["us_min"] == 9
    assert saved["measured_values"]["functional_observation"] == "No reset observed."


def test_result_evaluation_pass_case(client, sample_dut):
    plan_response = client.post("/api/plan/generate", json={"dut_id": sample_dut["id"]})
    test_id = plan_response.get_json()["data"]["tests"][0]["id"]

    response = client.post(
        "/api/result",
        json={
            "test_id": test_id,
            "dut_id": sample_dut["id"],
            "measured_values": {
                "us_min": 9,
                "us_max": 16,
                "duration": 2,
                "functional_observation": "Normal operation observed.",
            },
        },
    )

    assert response.status_code == 201
    saved = response.get_json()["data"]
    assert saved["evaluation_status"] == "PASS"
    assert saved["evaluation_score"] == 100
    assert saved["result"] == "Pass"


def test_result_evaluation_fail_case(client, sample_dut):
    plan_response = client.post("/api/plan/generate", json={"dut_id": sample_dut["id"]})
    test_id = plan_response.get_json()["data"]["tests"][0]["id"]

    response = client.post(
        "/api/result",
        json={
            "test_id": test_id,
            "dut_id": sample_dut["id"],
            "measured_values": {
                "us_min": 3,
                "us_max": 30,
                "duration": 0,
                "functional_observation": "Reset observed.",
            },
        },
    )

    assert response.status_code == 201
    saved = response.get_json()["data"]
    assert saved["evaluation_status"] == "FAIL"
    assert saved["evaluation_score"] < 80
    assert saved["evaluation_details"]["failed_rules"]


def test_result_evaluate_endpoint(client, sample_dut):
    plan_response = client.post("/api/plan/generate", json={"dut_id": sample_dut["id"]})
    test_id = plan_response.get_json()["data"]["tests"][0]["id"]

    response = client.post(
        f"/api/result/evaluate/{test_id}",
        json={
            "measured_values": {
                "us_min": 9,
                "us_max": 16,
                "duration": 2,
                "functional_observation": "Normal operation observed.",
            }
        },
    )

    assert response.status_code == 200
    body = response.get_json()["data"]
    assert body["status"] == "PASS"
    assert body["score"] == 100


def test_result_evaluation_no_schema_case(app, client, sample_dut):
    test_id = app.db.insert(
        """
        INSERT INTO tests (
            dut_id, test_name, standard_reference, category, status,
            duration_hours, required_equipment, acceptance_criteria, severity_level
        ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)
        """,
        (
            sample_dut["id"],
            "Legacy Result Test",
            "Legacy reference",
            "electrical",
            "Mandatory",
            1,
            "[]",
            "Legacy criteria",
            "I",
        ),
    )

    response = client.post(
        f"/api/result/evaluate/{test_id}",
        json={"measured_values": {"anything": "value"}},
    )

    assert response.status_code == 200
    assert response.get_json()["data"]["status"] == "NOT EVALUATED"


def test_create_and_list_equipment(client):
    response = client.post(
        "/api/equipment",
        json={
            "equipment_no": "EQ-001",
            "kind_of_equipment": "DC Power Supply",
            "model": "PSU-500",
            "type": "Supply",
            "manufacturer": "Acme Instruments",
            "serial_no": "SN-001",
            "next_calibration_date": (date.today() + timedelta(days=90)).isoformat(),
            "using_status": "available",
            "location": "Electrical Lab",
        },
    )

    assert response.status_code == 201
    item = response.get_json()["data"]
    assert item["equipment_no"] == "EQ-001"
    assert item["calibration_status"] == "valid"

    list_response = client.get("/api/equipment")

    assert list_response.status_code == 200
    items = list_response.get_json()["data"]["equipment"]
    assert any(equipment["equipment_no"] == "EQ-001" for equipment in items)


def test_delete_equipment(client):
    equipment = client.post(
        "/api/equipment",
        json={"equipment_no": "EQ-DELETE", "kind_of_equipment": "Power Analyzer"},
    ).get_json()["data"]

    delete_response = client.delete(f"/api/equipment/{equipment['id']}")
    get_response = client.get(f"/api/equipment/{equipment['id']}")

    assert delete_response.status_code == 200
    assert get_response.status_code == 404


def test_equipment_calibration_status_valid_due_soon_and_expired(client):
    dates = {
        "valid": date.today() + timedelta(days=45),
        "due_soon": date.today() + timedelta(days=10),
        "expired": date.today() - timedelta(days=1),
    }

    for expected_status, next_date in dates.items():
        response = client.post(
            "/api/equipment",
            json={
                "equipment_no": f"EQ-{expected_status}",
                "kind_of_equipment": "Multimeter",
                "next_calibration_date": next_date.isoformat(),
            },
        )
        assert response.status_code == 201
        assert response.get_json()["data"]["calibration_status"] == expected_status


def test_link_list_and_unlink_equipment_to_test(client, sample_dut):
    plan_response = client.post("/api/plan/generate", json={"dut_id": sample_dut["id"]})
    test_id = plan_response.get_json()["data"]["tests"][0]["id"]
    equipment = client.post(
        "/api/equipment",
        json={
            "equipment_no": "EQ-LINK",
            "kind_of_equipment": "Oscilloscope",
            "next_calibration_date": (date.today() + timedelta(days=90)).isoformat(),
        },
    ).get_json()["data"]

    link_response = client.post(
        "/api/equipment/link",
        json={
            "test_id": test_id,
            "equipment_id": equipment["id"],
            "usage_role": "Voltage monitoring",
        },
    )

    assert link_response.status_code == 201
    link = link_response.get_json()["data"]
    assert link["usage_role"] == "Voltage monitoring"
    assert link["equipment_no"] == "EQ-LINK"

    list_response = client.get(f"/api/equipment/test/{test_id}")
    assert list_response.status_code == 200
    assert list_response.get_json()["data"]["equipment"][0]["equipment_no"] == "EQ-LINK"

    unlink_response = client.delete(f"/api/equipment/link/{link['link_id']}")
    assert unlink_response.status_code == 200

    list_after_unlink = client.get(f"/api/equipment/test/{test_id}")
    assert list_after_unlink.get_json()["data"]["equipment"] == []


def test_record_form_includes_linked_equipment(client, sample_dut):
    plan_response = client.post("/api/plan/generate", json={"dut_id": sample_dut["id"]})
    test_id = plan_response.get_json()["data"]["tests"][0]["id"]
    equipment = client.post(
        "/api/equipment",
        json={
            "equipment_no": "EQ-DOCX",
            "kind_of_equipment": "Temperature Chamber",
            "model": "CH-1000",
            "type": "Chamber",
            "manufacturer": "Climate Systems",
            "serial_no": "CH-SN-1",
            "last_calibration_date": date.today().isoformat(),
            "next_calibration_date": (date.today() + timedelta(days=90)).isoformat(),
            "using_status": "available",
        },
    ).get_json()["data"]
    client.post(
        "/api/equipment/link",
        json={"test_id": test_id, "equipment_id": equipment["id"], "usage_role": "Thermal exposure"},
    )

    response = client.post(f"/api/record-form/generate/{test_id}")

    assert response.status_code == 200
    text = _docx_text(response.data)
    assert "EQ-DOCX" in text
    assert "Temperature Chamber" in text
    assert "CH-1000" in text


def test_technical_report_includes_linked_equipment(client, sample_dut):
    plan_response = client.post("/api/plan/generate", json={"dut_id": sample_dut["id"]})
    test_id = plan_response.get_json()["data"]["tests"][0]["id"]
    equipment = client.post(
        "/api/equipment",
        json={
            "equipment_no": "EQ-TR",
            "kind_of_equipment": "Data Logger",
            "model": "DL-20",
            "manufacturer": "Lab Systems",
            "next_calibration_date": (date.today() + timedelta(days=90)).isoformat(),
        },
    ).get_json()["data"]
    client.post(
        "/api/equipment/link",
        json={"test_id": test_id, "equipment_id": equipment["id"], "usage_role": "Measurement logging"},
    )

    response = client.post(
        f"/api/technical-report/generate/{sample_dut['id']}",
        json={"category": "all", "format": "docx"},
    )

    assert response.status_code == 200
    text = _docx_text(response.data)
    assert "EQ-TR" in text
    assert "Data Logger" in text


def test_record_form_endpoint_returns_docx_without_result(client, sample_dut):
    plan_response = client.post("/api/plan/generate", json={"dut_id": sample_dut["id"]})
    test_id = plan_response.get_json()["data"]["tests"][0]["id"]

    response = client.post(f"/api/record-form/generate/{test_id}")

    assert response.status_code == 200
    assert response.content_type.startswith(
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    assert len(response.data) > 1000


def test_attachment_upload_valid_file(client, sample_dut):
    plan_response = client.post("/api/plan/generate", json={"dut_id": sample_dut["id"]})
    test_id = plan_response.get_json()["data"]["tests"][0]["id"]

    response = client.post(
        "/api/attachments/upload",
        data={
            "dut_id": str(sample_dut["id"]),
            "test_id": str(test_id),
            "attachment_type": "measurement_file",
            "description": "Oscilloscope capture",
            "file": (BytesIO(b"voltage,current\n12,1.2\n"), "capture.csv"),
        },
        content_type="multipart/form-data",
    )

    assert response.status_code == 201
    attachment = response.get_json()["data"]
    assert attachment["original_filename"] == "capture.csv"
    assert attachment["attachment_type"] == "measurement_file"
    assert attachment["file_size"] > 0


def test_attachment_rejects_invalid_extension(client, sample_dut):
    plan_response = client.post("/api/plan/generate", json={"dut_id": sample_dut["id"]})
    test_id = plan_response.get_json()["data"]["tests"][0]["id"]

    response = client.post(
        "/api/attachments/upload",
        data={
            "dut_id": str(sample_dut["id"]),
            "test_id": str(test_id),
            "attachment_type": "raw_data",
            "file": (BytesIO(b"bad"), "script.exe"),
        },
        content_type="multipart/form-data",
    )

    assert response.status_code == 400
    assert "Invalid file extension" in response.get_json()["error"]


def test_attachment_list_by_test_and_dut(client, sample_dut):
    plan_response = client.post("/api/plan/generate", json={"dut_id": sample_dut["id"]})
    test_id = plan_response.get_json()["data"]["tests"][0]["id"]
    client.post(
        "/api/attachments/upload",
        data={
            "dut_id": str(sample_dut["id"]),
            "test_id": str(test_id),
            "attachment_type": "test_photo",
            "file": (BytesIO(b"fake image"), "photo.jpg"),
        },
        content_type="multipart/form-data",
    )

    by_test = client.get(f"/api/attachments/test/{test_id}")
    by_dut = client.get(f"/api/attachments/dut/{sample_dut['id']}")

    assert by_test.status_code == 200
    assert by_dut.status_code == 200
    assert len(by_test.get_json()["data"]["attachments"]) == 1
    assert len(by_dut.get_json()["data"]["attachments"]) == 1


def test_attachment_download_file(client, sample_dut):
    plan_response = client.post("/api/plan/generate", json={"dut_id": sample_dut["id"]})
    test_id = plan_response.get_json()["data"]["tests"][0]["id"]
    upload = client.post(
        "/api/attachments/upload",
        data={
            "dut_id": str(sample_dut["id"]),
            "test_id": str(test_id),
            "attachment_type": "temperature_humidity_log",
            "file": (BytesIO(b"temperature=23"), "log.txt"),
        },
        content_type="multipart/form-data",
    )
    attachment_id = upload.get_json()["data"]["id"]

    response = client.get(f"/api/attachments/download/{attachment_id}")

    assert response.status_code == 200
    assert response.data == b"temperature=23"


def test_attachment_delete_file(client, sample_dut):
    plan_response = client.post("/api/plan/generate", json={"dut_id": sample_dut["id"]})
    test_id = plan_response.get_json()["data"]["tests"][0]["id"]
    upload = client.post(
        "/api/attachments/upload",
        data={
            "dut_id": str(sample_dut["id"]),
            "test_id": str(test_id),
            "attachment_type": "supporting_document",
            "file": (BytesIO(b"note"), "note.txt"),
        },
        content_type="multipart/form-data",
    )
    attachment_id = upload.get_json()["data"]["id"]

    response = client.delete(f"/api/attachments/{attachment_id}")
    list_response = client.get(f"/api/attachments/test/{test_id}")

    assert response.status_code == 200
    assert list_response.get_json()["data"]["attachments"] == []


def test_record_form_lists_non_image_attachment(client, sample_dut):
    plan_response = client.post("/api/plan/generate", json={"dut_id": sample_dut["id"]})
    test_id = plan_response.get_json()["data"]["tests"][0]["id"]
    client.post(
        "/api/attachments/upload",
        data={
            "dut_id": str(sample_dut["id"]),
            "test_id": str(test_id),
            "attachment_type": "measurement_file",
            "description": "Recorded current log",
            "file": (BytesIO(b"current=1.2"), "current_log.txt"),
        },
        content_type="multipart/form-data",
    )

    response = client.post(f"/api/record-form/generate/{test_id}")

    assert response.status_code == 200
    text = _docx_text(response.data)
    assert "current_log.txt" in text
    assert "Recorded current log" in text
    assert "Table 1: DUT Identity and Test Requirement Information" in text
    assert "Table 4: Test Parameters" in text
    assert "Table 11: Signatures" in text
    assert "Required input" not in text
    assert "Mandatory recorded value" not in text


def test_technical_report_lists_non_image_attachment(client, sample_dut):
    plan_response = client.post("/api/plan/generate", json={"dut_id": sample_dut["id"]})
    test_id = plan_response.get_json()["data"]["tests"][0]["id"]
    client.post(
        "/api/attachments/upload",
        data={
            "dut_id": str(sample_dut["id"]),
            "test_id": str(test_id),
            "attachment_type": "raw_data",
            "description": "Raw voltage data",
            "file": (BytesIO(b"raw"), "raw_voltage.csv"),
        },
        content_type="multipart/form-data",
    )

    response = client.post(
        f"/api/technical-report/generate/{sample_dut['id']}",
        json={"category": "all", "format": "docx"},
    )

    assert response.status_code == 200
    text = _docx_text(response.data)
    assert "raw_voltage.csv" in text
    assert "Raw Data" in text
    assert "Test Standard" in text
    assert "Table 3.1: List of the Applied Tests" in text
    assert "Measured / Applied Values" in text
    assert "Evaluation Result" in text
    assert "Required input" not in text
    assert "Mandatory recorded value" not in text


def test_image_attachment_does_not_crash_document_generation(client, sample_dut):
    plan_response = client.post("/api/plan/generate", json={"dut_id": sample_dut["id"]})
    test_id = plan_response.get_json()["data"]["tests"][0]["id"]
    client.post(
        "/api/attachments/upload",
        data={
            "dut_id": str(sample_dut["id"]),
            "test_id": str(test_id),
            "attachment_type": "test_photo",
            "description": "Setup photo",
            "file": (BytesIO(b"not a real png"), "setup.png"),
        },
        content_type="multipart/form-data",
    )

    record_response = client.post(f"/api/record-form/generate/{test_id}")
    technical_response = client.post(
        f"/api/technical-report/generate/{sample_dut['id']}",
        json={"category": "all", "format": "docx"},
    )

    assert record_response.status_code == 200
    assert technical_response.status_code == 200
    assert "setup.png" in _docx_text(record_response.data)
    assert "setup.png" in _docx_text(technical_response.data)


def test_record_form_endpoint_returns_docx_with_result(client, sample_dut):
    plan_response = client.post("/api/plan/generate", json={"dut_id": sample_dut["id"]})
    test_id = plan_response.get_json()["data"]["tests"][0]["id"]
    client.post(
        "/api/result",
        json={
            "test_id": test_id,
            "dut_id": sample_dut["id"],
            "measured_values": {
                "us_min": 9,
                "us_max": 16,
                "duration": 2,
                "functional_observation": "Normal operation observed.",
            },
            "observations": "No anomaly observed.",
            "engineer_name": "Jane Engineer",
        },
    )

    response = client.post(f"/api/record-form/generate/{test_id}")

    assert response.status_code == 200
    assert len(response.data) > 1000
    _assert_no_turkish_docx_terms(response.data)


def test_record_form_endpoint_works_without_iso_catalog(app, client, sample_dut):
    test_id = app.db.insert(
        """
        INSERT INTO tests (
            dut_id, test_name, standard_reference, category, status,
            duration_hours, required_equipment, acceptance_criteria, severity_level
        ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)
        """,
        (
            sample_dut["id"],
            "Legacy Record Form Test",
            "Legacy reference",
            "electrical",
            "Mandatory",
            1,
            "[]",
            "Legacy criteria",
            "I",
        ),
    )

    response = client.post(f"/api/record-form/generate/{test_id}")

    assert response.status_code == 200
    assert len(response.data) > 1000


def test_technical_report_endpoint_returns_docx(client, sample_dut):
    client.post("/api/plan/generate", json={"dut_id": sample_dut["id"]})

    response = client.post(
        f"/api/technical-report/generate/{sample_dut['id']}",
        json={"category": "all", "format": "docx"},
    )

    assert response.status_code == 200
    assert response.content_type.startswith(
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    assert len(response.data) > 1000
    _assert_no_turkish_docx_terms(response.data)


def test_technical_report_endpoint_can_include_ai_summary(app, client, sample_dut):
    app.report_service.ai_client = FakeAIClient(
        {
            "executive_summary": "The selected ISO 16750 test scope was reviewed against recorded evidence.",
            "conformity_assessment": "The conformity statement is based on measured values and evaluation status.",
            "technical_comments": "No additional laboratory restriction was identified from the provided records.",
            "limitations": "This AI-assisted text does not replace the structured test tables.",
        }
    )
    client.post("/api/plan/generate", json={"dut_id": sample_dut["id"]})

    response = client.post(
        f"/api/technical-report/generate/{sample_dut['id']}",
        json={"category": "all", "format": "docx", "use_ai": True},
    )

    assert response.status_code == 200
    text = _docx_text(response.data)
    assert "AI-Assisted Technical Summary" in text
    assert "The selected ISO 16750 test scope was reviewed against recorded evidence." in text
    assert "This AI-assisted text does not replace the structured test tables." in text
    _assert_no_turkish_docx_terms(response.data)


def test_legacy_report_docx_export_has_no_turkish_terms(app, client, sample_dut):
    app.report_service.ai_client = FakeAIClient(
        {
            "sections": {
                "test_amaci": "The purpose is to document the laboratory test scope.",
                "test_kosullari": "The test was performed under recorded laboratory conditions.",
                "olcum_sonuclari": "Measured values were recorded in the result table.",
                "gozlemler": "No abnormal observation was recorded.",
                "kabul_degerlendirme": "Recorded values were compared with the acceptance criteria.",
                "sonuc": "The final decision is based on the recorded evidence.",
            }
        }
    )
    plan_response = client.post("/api/plan/generate", json={"dut_id": sample_dut["id"]})
    test_id = plan_response.get_json()["data"]["tests"][0]["id"]
    client.post(
        "/api/result",
        json={
            "test_id": test_id,
            "dut_id": sample_dut["id"],
            "measured_values": {
                "us_min": 9,
                "us_max": 16,
                "duration": 2,
                "functional_observation": "Normal operation observed.",
            },
            "observations": "No anomaly observed.",
            "engineer_name": "Jane Engineer",
        },
    )
    client.post(
        "/api/report/generate",
        json={"dut_id": sample_dut["id"], "test_id": test_id},
    )

    response = client.post(
        "/api/report/export/docx",
        json={"dut_id": sample_dut["id"], "test_id": test_id},
    )

    assert response.status_code == 200
    assert response.content_type.startswith(
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    assert len(response.data) > 1000
    _assert_no_turkish_docx_terms(response.data)


def test_legacy_report_docx_export_includes_uploaded_photo_reference(app, client, sample_dut):
    app.report_service.ai_client = FakeAIClient(
        {
            "sections": {
                "test_amaci": "The purpose is to document the laboratory test scope.",
                "test_kosullari": "The test was performed under recorded laboratory conditions.",
                "olcum_sonuclari": "Measured values were recorded in the result table.",
                "gozlemler": "No abnormal observation was recorded.",
                "kabul_degerlendirme": "Recorded values were compared with the acceptance criteria.",
                "sonuc": "The final decision is based on the recorded evidence.",
            }
        }
    )
    plan_response = client.post("/api/plan/generate", json={"dut_id": sample_dut["id"]})
    test_id = plan_response.get_json()["data"]["tests"][0]["id"]
    client.post(
        "/api/result",
        json={
            "test_id": test_id,
            "dut_id": sample_dut["id"],
            "measured_values": {
                "us_min": 9,
                "us_max": 16,
                "duration": 2,
                "functional_observation": "Normal operation observed.",
            },
        },
    )
    client.post(
        "/api/attachments/upload",
        data={
            "dut_id": str(sample_dut["id"]),
            "test_id": str(test_id),
            "attachment_type": "test_photo",
            "description": "Setup view",
            "file": (BytesIO(b"not a real image"), "setup_view.png"),
        },
        content_type="multipart/form-data",
    )
    client.post(
        "/api/report/generate",
        json={"dut_id": sample_dut["id"], "test_id": test_id},
    )

    response = client.post(
        "/api/report/export/docx",
        json={"dut_id": sample_dut["id"], "test_id": test_id},
    )

    assert response.status_code == 200
    text = _docx_text(response.data)
    assert "Attachments and Photo Evidence" in text
    assert "setup_view.png" in text
    assert "Setup view" in text


def test_technical_report_electrical_category_filters_tests(client, sample_dut):
    client.post("/api/plan/generate", json={"dut_id": sample_dut["id"]})

    response = client.post(
        f"/api/technical-report/generate/{sample_dut['id']}",
        json={"category": "electrical", "format": "docx"},
    )

    assert response.status_code == 200
    text = _docx_text(response.data)
    assert "DC Supply Voltage" in text
    assert "Low Temperature Operation" not in text
    assert "Chemical Resistance" not in text


def test_technical_report_environmental_category_filters_tests(client, sample_dut):
    client.post("/api/plan/generate", json={"dut_id": sample_dut["id"]})

    response = client.post(
        f"/api/technical-report/generate/{sample_dut['id']}",
        json={"category": "environmental", "format": "docx"},
    )

    assert response.status_code == 200
    text = _docx_text(response.data)
    assert "Low Temperature Operation" in text
    assert "DC Supply Voltage" not in text
    assert "Chemical Resistance" not in text


def test_technical_report_chemical_category_filters_tests(client, sample_dut):
    client.post("/api/plan/generate", json={"dut_id": sample_dut["id"]})

    response = client.post(
        f"/api/technical-report/generate/{sample_dut['id']}",
        json={"category": "chemical", "format": "docx"},
    )

    assert response.status_code == 200
    text = _docx_text(response.data)
    assert "Chemical Resistance" in text
    assert "DC Supply Voltage" not in text
    assert "Low Temperature Operation" not in text


def test_technical_report_missing_result_data_does_not_crash(client, sample_dut):
    client.post("/api/plan/generate", json={"dut_id": sample_dut["id"]})

    response = client.post(
        f"/api/technical-report/generate/{sample_dut['id']}",
        json={"category": "electrical", "format": "docx"},
    )

    assert response.status_code == 200
    assert len(response.data) > 1000


def _docx_text(data):
    document = Document(BytesIO(data))
    parts = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


def test_result_required_schema_validation(client, sample_dut):
    plan_response = client.post("/api/plan/generate", json={"dut_id": sample_dut["id"]})
    test_id = plan_response.get_json()["data"]["tests"][0]["id"]

    response = client.post(
        "/api/result",
        json={
            "test_id": test_id,
            "dut_id": sample_dut["id"],
            "result": "Pass",
            "measured_values": {"us_min": 9},
        },
    )

    assert response.status_code == 400
    assert "required" in response.get_json()["error"]


def test_result_number_schema_validation(client, sample_dut):
    plan_response = client.post("/api/plan/generate", json={"dut_id": sample_dut["id"]})
    test_id = plan_response.get_json()["data"]["tests"][0]["id"]

    response = client.post(
        "/api/result",
        json={
            "test_id": test_id,
            "dut_id": sample_dut["id"],
            "result": "Pass",
            "measured_values": {
                "us_min": "not-a-number",
                "us_max": 16,
                "duration": 2,
                "functional_observation": "No reset observed.",
            },
        },
    )

    assert response.status_code == 400
    assert "must be a number" in response.get_json()["error"]


def test_dashboard_summary(client, sample_dut):
    client.post("/api/plan/generate", json={"dut_id": sample_dut["id"]})

    response = client.get("/api/dashboard/summary")
    assert response.status_code == 200
    data = response.get_json()["data"]
    assert data["summary"]["total_duts"] == 1
    assert data["summary"]["total_tests"] > 0
    assert data["summary"]["pending_tests"] == data["summary"]["total_tests"]
    assert len(data["duts"]) == 1
    assert len(data["duts"][0]["tests"]) == data["summary"]["total_tests"]


def test_dashboard_overview_empty_database(client):
    response = client.get("/api/dashboard/overview")

    assert response.status_code == 200
    data = response.get_json()["data"]
    for key in (
        "summary",
        "tests_by_category",
        "tests_by_status",
        "results_by_evaluation_status",
        "equipment_calibration_summary",
        "workflow_progress",
        "recent_duts",
        "recent_test_results",
        "attention_required",
    ):
        assert key in data
    assert data["summary"]["total_duts"] == 0
    assert data["summary"]["total_tests"] == 0
    assert data["workflow_progress"]["record_forms_available"] == "Available on demand"
    assert data["attention_required"]["failed_tests"] == []


def test_dashboard_overview_with_sample_workflow(client, sample_dut):
    plan_response = client.post("/api/plan/generate", json={"dut_id": sample_dut["id"]})
    test_id = plan_response.get_json()["data"]["tests"][0]["id"]
    equipment = client.post(
        "/api/equipment",
        json={
            "equipment_no": "EQ-OVERVIEW",
            "kind_of_equipment": "DC Power Supply",
            "next_calibration_date": (date.today() - timedelta(days=1)).isoformat(),
        },
    ).get_json()["data"]
    client.post(
        "/api/equipment/link",
        json={"test_id": test_id, "equipment_id": equipment["id"], "usage_role": "Supply voltage"},
    )
    client.post(
        "/api/attachments/upload",
        data={
            "dut_id": str(sample_dut["id"]),
            "test_id": str(test_id),
            "attachment_type": "measurement_file",
            "file": (BytesIO(b"voltage,current\n12,1\n"), "overview.csv"),
        },
        content_type="multipart/form-data",
    )
    client.post(
        "/api/result",
        json={
            "test_id": test_id,
            "dut_id": sample_dut["id"],
            "measured_values": {
                "us_min": 9,
                "us_max": 16,
                "duration": 2,
                "functional_observation": "Normal operation observed.",
            },
        },
    )

    response = client.get("/api/dashboard/overview")

    assert response.status_code == 200
    data = response.get_json()["data"]
    assert data["summary"]["total_duts"] == 1
    assert data["summary"]["total_plan_items"] > 0
    assert data["summary"]["total_tests"] > 0
    assert data["summary"]["completed_results"] == 1
    assert data["summary"]["attachments"] == 1
    assert data["summary"]["linked_equipment_items"] == 1
    assert data["equipment_calibration_summary"]["expired"] == 1
    assert data["results_by_evaluation_status"]["PASS"] == 1
    assert data["recent_duts"][0]["name"] == sample_dut["name"]
    assert data["recent_test_results"][0]["evaluation_status"] == "PASS"
    assert data["attention_required"]["expired_equipment"][0]["equipment_no"] == "EQ-OVERVIEW"


def test_demo_seed_endpoint_creates_presentation_dataset(app, client):
    response = client.post("/api/demo/seed", json={"reset_demo_data": True})

    assert response.status_code == 201
    data = response.get_json()["data"]
    assert data["created"] is True
    assert data["dut"] == "Stage V Instrument Cluster"
    assert data["tests"] == 5
    assert data["plan_items"] == 5
    assert data["results"] == 5
    assert data["equipment"] == 3
    assert data["equipment_links"] >= 4
    assert data["attachments"] == 2

    dut = app.db.query_one("SELECT * FROM duts WHERE id = ?", (data["dut_id"],))
    assert "DEMO_DATA" in dut["notes"]
    assert dut["name"] == "Stage V Instrument Cluster"
    assert dut["part_number"] == "BC400 TT/B"

    tests = app.db.query("SELECT * FROM tests WHERE dut_id = ?", (data["dut_id"],))
    assert {test["test_name"] for test in tests} >= {
        "DC Supply Voltage",
        "Superimposed Alternating Voltage",
        "Overvoltage",
        "Low Temperature Operation",
        "Chemical Resistance / Chemical Loads",
    }
    assert app.db.query_one("SELECT COUNT(*) AS count FROM test_plan_items WHERE dut_id = ?", (data["dut_id"],))["count"] == 5
    assert app.db.query_one("SELECT COUNT(*) AS count FROM results WHERE dut_id = ?", (data["dut_id"],))["count"] == 5
    result_rows = app.db.query(
        """
        SELECT t.test_name, r.measured_values
        FROM results r
        JOIN tests t ON t.id = r.test_id
        WHERE r.dut_id = ?
        """,
        (data["dut_id"],),
    )
    assert {row["test_name"] for row in result_rows} >= {
        "DC Supply Voltage",
        "Superimposed Alternating Voltage",
        "Overvoltage",
        "Low Temperature Operation",
        "Chemical Resistance / Chemical Loads",
    }
    assert all(row["measured_values"] for row in result_rows)
    assert app.db.query_one("SELECT COUNT(*) AS count FROM equipment WHERE notes LIKE ?", ("%DEMO_DATA%",))["count"] == 3
    assert app.db.query_one("SELECT COUNT(*) AS count FROM test_equipment")["count"] >= 4
    assert app.db.query_one("SELECT COUNT(*) AS count FROM test_attachments WHERE dut_id = ?", (data["dut_id"],))["count"] == 2


def test_demo_status_endpoint_when_demo_data_is_missing(client):
    response = client.get("/api/demo/status")

    assert response.status_code == 200
    data = response.get_json()["data"]
    assert data["demo_data_loaded"] is False
    assert data["demo_dut_id"] is None
    assert data["demo_project_no"] is None
    assert data["demo_test_count"] == 0
    assert data["demo_result_count"] == 0
    assert data["demo_attachment_count"] == 0
    assert data["demo_equipment_count"] == 0
    assert data["recommended_start_route"] == "/dashboard"
    assert "Demo DUT" in data["missing_demo_items"]


def test_demo_status_endpoint_when_demo_data_exists(client):
    seed = client.post("/api/demo/seed", json={"reset_demo_data": True}).get_json()["data"]

    response = client.get("/api/demo/status")

    assert response.status_code == 200
    data = response.get_json()["data"]
    assert data["demo_data_loaded"] is True
    assert data["demo_dut_id"] == seed["dut_id"]
    assert data["demo_project_no"] == "DEMO-ISO16750-STAGEV"
    assert data["demo_test_count"] == 5
    assert data["demo_result_count"] == 5
    assert data["demo_attachment_count"] == 2
    assert data["demo_equipment_count"] == 3
    assert data["first_demo_test_id"] is not None
    assert data["first_demo_result_test_id"] is not None
    assert data["recommended_start_route"] == f"/plan/{seed['dut_id']}"
    assert data["missing_demo_items"] == []


def test_demo_seed_endpoint_does_not_duplicate_without_reset(app, client):
    first = client.post("/api/demo/seed", json={"reset_demo_data": True}).get_json()["data"]
    second_response = client.post("/api/demo/seed", json={"reset_demo_data": False})
    second = second_response.get_json()["data"]

    assert second_response.status_code == 200
    assert second["created"] is False
    assert second["dut_id"] == first["dut_id"]
    assert app.db.query_one("SELECT COUNT(*) AS count FROM duts WHERE notes LIKE ?", ("%DEMO_DATA%",))["count"] == 1
    assert app.db.query_one("SELECT COUNT(*) AS count FROM tests WHERE dut_id = ?", (first["dut_id"],))["count"] == 5


def test_demo_seed_reset_recreates_demo_dataset(app, client):
    first = client.post("/api/demo/seed", json={"reset_demo_data": True}).get_json()["data"]
    second = client.post("/api/demo/seed", json={"reset_demo_data": True}).get_json()["data"]

    assert second["created"] is True
    assert second["dut_id"] != first["dut_id"]
    assert app.db.query_one("SELECT COUNT(*) AS count FROM duts WHERE notes LIKE ?", ("%DEMO_DATA%",))["count"] == 1
    assert app.db.query_one("SELECT COUNT(*) AS count FROM equipment WHERE notes LIKE ?", ("%DEMO_DATA%",))["count"] == 3


def test_dashboard_overview_reflects_demo_seed(client):
    client.post("/api/demo/seed", json={"reset_demo_data": True})

    response = client.get("/api/dashboard/overview")

    assert response.status_code == 200
    data = response.get_json()["data"]
    assert data["summary"]["total_duts"] == 1
    assert data["summary"]["total_plan_items"] == 5
    assert data["summary"]["total_tests"] == 5
    assert data["summary"]["completed_results"] == 5
    assert data["summary"]["attachments"] == 2
    assert data["summary"]["linked_equipment_items"] >= 4
    assert data["results_by_evaluation_status"]["PASS"] >= 1
    assert data["results_by_evaluation_status"]["FAIL"] >= 1
    assert data["results_by_evaluation_status"]["PASS"] + data["results_by_evaluation_status"]["FAIL"] == 5
    assert data["equipment_calibration_summary"]["valid"] == 1
    assert data["equipment_calibration_summary"]["due_soon"] == 1
    assert data["equipment_calibration_summary"]["expired"] == 1
