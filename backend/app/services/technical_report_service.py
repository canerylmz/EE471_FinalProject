"""DOCX generator for ISO 16750:2023 final technical reports."""

import io
import json
import re
from collections import OrderedDict
from datetime import date

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from ..utils.attachment_docx import (
    ATTACHMENT_TYPE_LABELS,
    add_attachment_table,
    add_image_thumbnail_safe,
    group_attachments_by_type,
    is_image_attachment,
)


DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
CATEGORY_PREFIXES = {"electrical": "EL", "environmental": "ENV", "chemical": "CH", "all": "ISO"}
VALID_CATEGORIES = set(CATEGORY_PREFIXES)
ISO_PART_FOR_CATEGORY = {
    "electrical": "ISO 16750-2:2023",
    "environmental": "ISO 16750-4:2023",
    "chemical": "ISO 16750-5:2023",
}
HEADER_BLUE = RGBColor(0x1F, 0x4E, 0x79)
TOC_BLUE = RGBColor(0x2F, 0x54, 0x96)
HEADER_FILL = "D9D9D9"

CHEMICAL_AGENT_IDS = (
    ("diesel", "AA"),
    ("bio diesel", "AB"),
    ("biodiesel", "AB"),
    ("antifreeze", "CC"),
    ("coolant", "CC"),
)

PHYSICAL_CHECK_ROWS = (
    (
        "Non-Metallic Parts",
        "Does any part of the assembly fall apart?",
        ("fall apart", "falls apart", "fell apart"),
    ),
    (
        "Non-Metallic Parts",
        "Is there crazing or swelling?",
        ("craz", "swell"),
    ),
    (
        "Non-Metallic Parts",
        "Is there any leeching of fluids (fluid becomes a permanent part of the surface)?",
        ("leech",),
    ),
    (
        "Non-Metallic Parts",
        "Are there any failures of seals or gaskets?",
        ("seal", "gasket"),
    ),
    (
        "Non-Metallic Parts",
        "Did any paint, stickers, or labels come off?",
        ("paint", "sticker", "label"),
    ),
    (
        "Non-Metallic Parts",
        "Was there any corrosion?",
        ("corrosion", "corroded"),
    ),
    (
        "Non-Metallic Parts",
        "Was there any melting or decomposition of the non-metallic parts?",
        ("melt", "decompos"),
    ),
)


class TechnicalReportService:
    """Builds formal ISO 16750:2023 final technical report DOCX files."""

    def __init__(self, organization_name="TestForge Test Laboratory", organization_address="-"):
        self.organization_name = organization_name
        self.organization_address = organization_address

    def build_docx(self, dut, test_entries, category, ai_sections=None):
        document = Document()
        report_no = self.report_number(dut, category)
        self._configure(document, report_no)

        self._cover_page(document, dut, test_entries, category, report_no)
        document.add_page_break()
        self._contents(document, test_entries, category)
        document.add_page_break()

        if ai_sections:
            self._ai_summary(document, ai_sections)
            self._page_break(document)

        self._sample_identification(document, dut, ai_sections)
        self._functional_status_classification(document)
        self._classification_by_mounting_location(document)
        self._operating_modes(document)
        self._parameter_check_levels(document)
        self._page_break(document)
        self._applied_tests(document, test_entries)
        self._test_requirements(document, dut)
        self._measurement_uncertainty(document)
        self._decision_rule(document)
        self._page_break(document)
        self._tests_section(document, test_entries, category)
        self._comments(document, test_entries, ai_sections)
        self._page_break(document)
        self._attachments(document, test_entries, category)

        buffer = io.BytesIO()
        document.save(buffer)
        buffer.seek(0)
        return buffer

    def report_number(self, dut, category):
        prefix = CATEGORY_PREFIXES.get(category, "ISO")
        project = dut.get("project") or f"DUT{dut.get('id')}"
        safe_project = "".join(c if c.isalnum() or c in ("-", "_") else "-" for c in str(project))
        return f"{prefix}-{safe_project}-ISO16750"

    # ------------------------------------------------------------------
    # Page setup, header, footer
    # ------------------------------------------------------------------

    def _configure(self, document, report_no):
        section = document.sections[0]
        section.page_width = Inches(8.27)
        section.page_height = Inches(11.69)
        section.top_margin = Inches(1560 / 1440)
        section.right_margin = Inches(1134 / 1440)
        section.bottom_margin = Inches(929 / 1440)
        section.left_margin = Inches(1134 / 1440)
        section.header_distance = Inches(680 / 1440)
        section.footer_distance = Inches(510 / 1440)

        style = document.styles["Normal"]
        style.font.name = "Calibri"
        style.font.size = Pt(12)

        header = section.header.paragraphs[0]
        header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = header.add_run(self.organization_name)
        run.bold = True
        run.font.size = Pt(9)

        footer = section.footer
        disclaimer = footer.paragraphs[0]
        disclaimer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = disclaimer.add_run(
            "This technical report can only be quoted as the whole part. It must be certified "
            "officially for all advertising purposes. This report is the result of only the "
            "tested sample and not a general assessment for the serial manufacturing products."
        )
        run.italic = True
        run.font.size = Pt(7)

        info = footer.add_paragraph()
        info.paragraph_format.tab_stops.add_tab_stop(Inches(3.5), WD_TAB_ALIGNMENT.CENTER)
        info.paragraph_format.tab_stops.add_tab_stop(Inches(7.0), WD_TAB_ALIGNMENT.RIGHT)
        info.add_run(f"Technical Report No.: {report_no}").font.size = Pt(7)
        info.add_run(f"\t{self.organization_name}, {self.organization_address}").font.size = Pt(7)
        info.add_run("\tDate of release: ").font.size = Pt(7)
        info.add_run(date.today().isoformat()).font.size = Pt(7)
        info.add_run("  Page: ").font.size = Pt(7)
        self._add_page_field(info, "PAGE")
        info.add_run(" / ").font.size = Pt(7)
        self._add_page_field(info, "NUMPAGES")

    # ------------------------------------------------------------------
    # Cover page
    # ------------------------------------------------------------------

    def _cover_page(self, document, dut, entries, category, report_no):
        label = document.add_paragraph()
        label.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = label.add_run("Technical Report No :")
        run.bold = True
        run.italic = True
        run.font.size = Pt(12)

        number = document.add_paragraph()
        number.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = number.add_run(report_no)
        run.bold = True
        run.font.size = Pt(16)
        run.font.color.rgb = HEADER_BLUE

        subtitle = document.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle.add_run(self._category_title(category).upper()).bold = True

        table = document.add_table(rows=0, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        self._style_table(table, header=False)
        for label_text, value, shaded in (
            ("Client", dut.get("customer"), True),
            ("Manufacturing Location", dut.get("manufacturer"), True),
            ("Test Sample / Models/Versions", dut.get("part_number") or dut.get("name"), True),
            ("Testing Location", "-", True),
            ("Test Standard", self._test_standard(entries, category), True),
            ("Purpose of Test", "Production Sample", True),
            ("Date of Receipt of Test Item", "-", True),
            ("Date(s) of Performance of Tests", self._test_dates(entries), True),
            ("Results", self._overall_result(entries), True),
            ("Date of Issue", date.today().isoformat(), True),
        ):
            self._cover_row(table, label_text, value)
        self._spacer(document)

    def _cover_row(self, table, label, value):
        row = table.add_row().cells
        row[0].text = ""
        run = row[0].paragraphs[0].add_run(label)
        run.italic = True
        self._set_cell_shading(row[0], HEADER_FILL)
        row[1].text = self._display(value)

    # ------------------------------------------------------------------
    # Contents
    # ------------------------------------------------------------------

    def _contents(self, document, entries, category):
        heading = document.add_paragraph()
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = heading.add_run("CONTENTS")
        run.bold = True
        run.font.color.rgb = TOC_BLUE

        rows = [
            ("AI.", "AI-Assisted Technical Summary"),
            ("1.", "Identification of the Test Sample"),
            ("1.1.", "Function"),
            ("1.2.", "Technical Information"),
            ("1.3.", "Test Configurations"),
            ("2.", "Classification of the Functional Status"),
            ("2.1.", "Classification by Mounting Location"),
            ("2.2.", "Operating Modes"),
            ("2.3.", "Parameter Check Levels"),
            ("3.", "List of the Applied Tests"),
            ("4.", "Test and Requirements"),
            ("5.", "Statement of Measurement Uncertainty"),
            ("6.", "Decision Rule & Conformity Assessment"),
            ("7.", "Tests"),
        ]
        for index, entry in enumerate(entries, start=1):
            rows.append((f"7.{index}.", entry["test"].get("test_name") or "Test"))
        rows.append(("8.", "Comments"))
        rows.append(("I.", "Attachment I: Photo Documentation"))
        rows.append(("II.", "Attachment II: Functional Checks for the DuT"))
        rows.append(("III.", f"Attachment III: {self._attachment_iii_title(entries)}"))

        for number, title in rows:
            self._contents_row(document, number, title)

    @staticmethod
    def _contents_row(document, number, title):
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.tab_stops.add_tab_stop(
            Inches(7.05),
            WD_TAB_ALIGNMENT.RIGHT,
            WD_TAB_LEADER.DOTS,
        )
        paragraph.paragraph_format.space_after = Pt(2)
        paragraph.add_run(f"{number}  {title}")
        paragraph.add_run("\t-")

    # ------------------------------------------------------------------
    # AI-assisted summary
    # ------------------------------------------------------------------

    def _ai_summary(self, document, ai_sections):
        self._heading(document, "AI-Assisted Technical Summary", 1)
        paragraph = document.add_paragraph(
            "This section was generated by the configured AI service from recorded DUT, test, "
            "measurement, equipment, attachment, and evaluation data. The structured tables in "
            "the following sections remain the authoritative test record."
        )
        paragraph.paragraph_format.space_after = Pt(6)

        self._table_caption(document, "AI Summary")
        table = document.add_table(rows=0, cols=2)
        self._style_table(table)
        for label, key in (
            ("Executive Summary", "executive_summary"),
            ("Conformity Assessment", "conformity_assessment"),
            ("Technical Comments", "technical_comments"),
            ("Limitations", "limitations"),
        ):
            self._field_row(table, label, ai_sections.get(key))
        self._spacer(document)

    # ------------------------------------------------------------------
    # Section 1: Identification of the Test Sample
    # ------------------------------------------------------------------

    def _sample_identification(self, document, dut, ai_sections):
        self._heading(document, "1. Identification of the Test Sample", 1)
        self._heading(document, "1.1 Function", 2)
        function_text = (ai_sections or {}).get("function_description") or (
            "Automotive electrical/electronic component under ISO 16750:2023 verification. "
            "The test scope was selected from the ISO 16750 catalog based on the recorded DUT metadata."
        )
        document.add_paragraph(function_text)
        self._spacer(document)

        self._heading(document, "1.2 Technical Information", 2)
        self._table_caption(document, "Table 1.2.1: Technical Information of the DuT")
        table = document.add_table(rows=0, cols=2)
        self._style_table(table)
        for label, value in (
            ("Product Description", dut.get("name")),
            ("Manufacturer Part Nr", dut.get("part_number")),
            ("Working Voltage", dut.get("nominal_voltage")),
            ("Operating Temperature", self._temperature_range(dut)),
            ("Mounting Location", dut.get("mounting_location")),
            ("Operating Current", "-"),
            ("IP Code", dut.get("ip_class")),
            ("Mass ( gr )", "-"),
            ("Dimensions ( mm )", "-"),
        ):
            self._field_row(table, label, value)
        self._spacer(document)

        self._heading(document, "1.3 Test Configurations", 2)
        document.add_paragraph(
            "The following configuration markings are used where the DuT was connected to a "
            "load/function simulation board during testing."
        )
        self._table_caption(document, "Table 1.3.1: Test Configurations")
        table = document.add_table(rows=1, cols=3)
        self._style_table(table)
        self._header(table.rows[0], ["Marking in Table", "Meaning", "Description"])
        for row in (
            ("TC1", "Direct connection without simulation board", "Representative DuT harness and supply interface"),
            ("TC2", "Single simulation board", "One DuT connected to project-specific load simulation"),
            ("TC3", "Multiple DuT configuration", "Multiple samples operated under the same exposure profile"),
        ):
            self._table_row(table, row)
        self._spacer(document)

    # ------------------------------------------------------------------
    # Section 2: Classification of the Functional Status
    # ------------------------------------------------------------------

    def _functional_status_classification(self, document):
        self._heading(document, "2. Classification of the Functional Status", 1)
        intro = document.add_paragraph()
        run = intro.add_run(
            "In this specification to classify the functional status for one, some or all the "
            "functions managed by the component (context dependent, meaning: for each applicable "
            "test the acceptance criteria for one, some or all the functions, will be selected "
            "picking one definition below), we will use the following definitions, as per "
            "ISO 16750-1:2023:"
        )
        run.italic = True

        classes = (
            (
                "Class A:",
                "All functions of the device/system perform as designed during the test (for the "
                "periods with active operating mode) and after the test.",
            ),
            (
                "Class B:",
                "All functions of the device/system perform as designed during the test. However, "
                "one or more of them may go beyond the specified tolerance. All functions "
                "automatically return to within normal limits after the test. Memory functions "
                "shall remain in class A. The vehicle manufacturer specifies which functions of the "
                "DUT shall perform as designed during the test and which functions can be beyond "
                "the specified tolerance.",
            ),
            (
                "Class C:",
                "One or more functions of a device/system do not perform as designed during the "
                "test, but automatically return to normal operation after the test.",
            ),
            (
                "Class D:",
                "One or more functions of a device/system do not perform as designed during the "
                "test and do not return to normal operation after the test until the device/system "
                "is reset by a simple operator/user action.",
            ),
            (
                "Class E:",
                "One or more functions of a device/system do not perform as designed during and "
                "after the test and cannot be returned to proper operation without repairing or "
                "replacing the device/system.",
            ),
        )
        for title, body in classes:
            paragraph = document.add_paragraph(style="List Bullet")
            run = paragraph.add_run(f"{title} ")
            run.bold = True
            paragraph.add_run(body)

        note = document.add_paragraph()
        run = note.add_run(
            "Note: Unwanted operations of the DUT are not allowed in any of the above classes "
            "(ISO 16750-1:2023, Clause 6.1)."
        )
        run.italic = True
        self._spacer(document)

    def _classification_by_mounting_location(self, document):
        self._heading(document, "2.1 Classification by Mounting Location", 2)
        intro = document.add_paragraph()
        intro.add_run(
            "Mounting locations are classified per ISO 16750-1:2023, Clause 4, as summarized in "
            "Table 2.1.1. The DuT mounting location for this report is specified in Section 1.2 "
            "(Technical Information)."
        ).bold = True

        self._table_caption(document, "Table 2.1.1: Classification by Mounting Location")
        table = document.add_table(rows=1, cols=3)
        self._style_table(table)
        self._header(table.rows[0], ["Location", "Description", "Sub-locations"])
        rows = (
            (
                "Engine/electric motor compartment",
                "Mounted to body, frame, on/in engine, on transmission",
                "front end upper/lower; higher/lower than members; on engine top/bottom; in "
                "engine; on/in transmission/gearbox",
            ),
            (
                "Passenger compartment",
                "Mounted inside",
                "without special req.; exposed to solar radiation; exposed to radiated heat",
            ),
            (
                "Luggage compartment / load compartment",
                "Mounted inside",
                "-",
            ),
            (
                "Mounting on exterior / in cavities",
                "External mounting",
                "to body top/side/bottom; underbody/wheel housing; door inside/outside; engine "
                "compartment cover; in cavities open to interior/exterior",
            ),
            (
                "Other mounting location",
                "Special locations (e.g. exhaust)",
                "As agreed between customer and supplier",
            ),
        )
        for location, description, sub in rows:
            row = table.add_row().cells
            row[0].paragraphs[0].add_run(location).bold = True
            row[1].text = description
            row[2].text = sub
        self._spacer(document)

    def _operating_modes(self, document):
        self._heading(document, "2.2 Operating Modes", 2)
        document.add_paragraph().add_run(
            "Operating modes are defined per ISO 16750-1:2023, Clause 5, as summarized in Table "
            "2.2.1."
        ).bold = True

        self._table_caption(document, "Table 2.2.1: Operating Modes for 12/24 V DUT")
        table = document.add_table(rows=1, cols=4)
        self._style_table(table)
        self._header(table.rows[0], ["Operating Mode", "Wire Harness Connected", "Supply Voltage", "Load Condition"])
        rows = (
            ("1.1", "No", "No applied voltage", "Not applicable"),
            ("1.2", "Yes", "No applied voltage", "Deactivated"),
            ("2.1", "Yes", "UB", "Sleep mode"),
            ("2.2", "Yes", "UB", "Typical mode"),
            ("2.3", "Yes", "UB", "Minimum load"),
            ("2.4", "Yes", "UB", "Maximum load"),
            ("3.1", "Yes", "UA", "Deactivated"),
            ("3.2", "Yes", "UA", "Typical mode (stand-by for ECUs with µP)"),
            ("3.3", "Yes", "UA", "Minimum load"),
            ("3.4", "Yes", "UA", "Maximum load"),
            ("4.1", "Yes", "UA + auxiliary device", "Deactivated"),
            ("4.2", "Yes", "UA + auxiliary device", "Typical mode"),
            ("4.3", "Yes", "UA + auxiliary device", "Minimum load"),
            ("4.4", "Yes", "UA + auxiliary device", "Maximum load"),
        )
        for om, harness, voltage, load in rows:
            row = table.add_row().cells
            row[0].paragraphs[0].add_run(om).bold = True
            row[1].text = harness
            row[2].text = voltage
            row[3].text = load

        footnote = document.add_paragraph()
        run = footnote.add_run(
            "UB = battery voltage (engine off); UA = charging voltage (engine running). For "
            "specific test cases, other voltage profiles apply (see ISO 16750-2)."
        )
        run.italic = True
        run.font.size = Pt(9)
        self._spacer(document)

    def _parameter_check_levels(self, document):
        self._heading(document, "2.3 Parameter Check Levels", 2)
        document.add_paragraph(
            "Parameter check and physical analysis requirements are defined per ISO 16750-1:2023, "
            "Clauses 7.6 and 7.7."
        )
        self._table_caption(document, "Table 2.3.1: Monitoring Requirement Levels")
        table = document.add_table(rows=1, cols=2)
        self._style_table(table)
        self._header(table.rows[0], ["Code", "Definition"])
        for code, definition in (
            ("PC0", "No parameter check required during or after test"),
            ("PC1", "Parameter check before and after test (key parameters as defined in component specification)"),
            ("PC2", "Continuous parameter monitoring during test and parameter check before and after test"),
        ):
            row = table.add_row().cells
            row[0].paragraphs[0].add_run(code).bold = True
            row[1].text = definition

        note = document.add_paragraph()
        run = note.add_run(
            "Note: Physical analysis / visual inspection (per ISO 16750-1:2023, Clause 7.7 and "
            "EN 13018) shall be carried out on at least one DUT after each test sequence."
        )
        run.italic = True
        self._spacer(document)

    # ------------------------------------------------------------------
    # Section 3: List of the Applied Tests
    # ------------------------------------------------------------------

    def _applied_tests(self, document, entries):
        self._heading(document, "3. List of the Applied Tests", 1)
        self._table_caption(document, "Table 3.1: List of the Applied Tests")
        table = document.add_table(rows=1, cols=7)
        self._style_table(table)
        self._header(
            table.rows[0],
            ["Test No.", "ISO 16750 Clause", "Name of the Test", "Functional Status Class", "Sample No.", "Deviations", "Result"],
        )
        for index, entry in enumerate(entries, start=1):
            test = entry["test"]
            result = entry.get("result") or {}
            self._table_row(
                table,
                (
                    test.get("planned_test_no") or f"7.{index}",
                    self._clause_reference(test),
                    test.get("test_name"),
                    test.get("functional_status"),
                    test.get("sample_size"),
                    self._deviation(result),
                    self._result_label(result),
                ),
            )
            sub = table.add_row().cells
            sub_text = f"OM {self._display(test.get('operating_mode'))} / {self._parameter_check_level(test)}"
            merged = sub[0].merge(sub[6])
            paragraph = merged.paragraphs[0]
            run = paragraph.add_run(sub_text)
            run.italic = True
            run.font.size = Pt(8)

        self._spacer(document)
        self._table_caption(document, "Legend")
        legend = document.add_paragraph()
        for line in (
            "N/A: Test is not applicable for the DuT",
            "ND: Test is not demanded by the manufacturer / customer",
            "Pass: Test object meets the requirement",
            "Fail: Test object does not meet the requirement",
        ):
            legend.add_run(line + "\n")

        self._spacer(document)
        self._table_caption(document, "Table 3.2: Recorded Deviations")
        deviations = document.add_table(rows=1, cols=2)
        self._style_table(deviations)
        self._header(deviations.rows[0], ["Deviation No.", "Description of the Deviation"])
        found = False
        for index, entry in enumerate(entries, start=1):
            result = entry.get("result") or {}
            if result.get("has_deviation"):
                found = True
                self._table_row(deviations, (f"No.:{index}", result.get("deviation_description")))
        if not found:
            self._table_row(deviations, ("No deviation", "No deviation has been recorded."))
        self._spacer(document)

    # ------------------------------------------------------------------
    # Section 4: Test and Requirements
    # ------------------------------------------------------------------

    def _test_requirements(self, document, dut):
        self._heading(document, "4. Test and Requirements", 1)
        self._table_caption(document, "Table 4.1: General Test Condition Tolerances (ISO 16750-2:2023, Clause 4.1)")
        table = document.add_table(rows=1, cols=2)
        self._style_table(table)
        self._header(table.rows[0], ["Parameter", "Tolerance"])
        for row in (
            ("Frequency and time", "± 5 %"),
            ("Voltages", "± 0,2 V"),
            ("Currents", "± 2 %"),
            ("Inductance", "± 10 %"),
            ("Resistance", "± 10 %"),
        ):
            self._table_row(table, row)
        self._spacer(document)

        document.add_paragraph(
            "For physical/environmental measurements, the following laboratory instrument "
            "tolerances apply:"
        )
        self._table_caption(document, "Table 4.2: Laboratory Instrument Tolerances")
        table = document.add_table(rows=1, cols=2)
        self._style_table(table)
        self._header(table.rows[0], ["Parameter", "Tolerance"])
        for row in (
            ("Temperature", "± 1 °C"),
            ("Relative Humidity", "± 6 %"),
            ("Pressure (Range)", "± 1,5 %"),
            ("Length, Distance", "± 1 %"),
            ("Mass, Weight, Acceleration", "± 1 %"),
            ("Volume", "± 5 %"),
        ):
            self._table_row(table, row)
        self._spacer(document)

        self._heading(document, "Supply Voltages", 2)
        nominal = (dut.get("nominal_voltage") or "").lower()
        show_12v = "24" not in nominal
        show_24v = "24" in nominal or "12" not in nominal
        if show_12v:
            document.add_paragraph(
                "For 12 V nominal systems (per ISO 16750-1:2023, Table 4 and ISO 16750-2:2023):"
            )
            for line in (
                "UA = 14 ± 0,2 V (engine running / charging voltage)",
                "UB = 12 ± 0,2 V (engine off / battery voltage)",
                "USmin/USmax per ISO 16750-2:2023, Table 3 (Code A: USmin = 9 V, USmax = 16 V; "
                "Code B: USmin = 6 V, USmax = 16 V; etc.)",
            ):
                document.add_paragraph(line, style="List Bullet")
        if show_24v:
            document.add_paragraph(
                "For 24 V nominal systems (per ISO 16750-1:2023, Table 4 and ISO 16750-2:2023):"
            )
            for line in (
                "UA = 28 ± 0,2 V (engine running)",
                "UB = 24 ± 0,2 V (engine off)",
                "USmin/USmax per ISO 16750-2:2023, Table 4 (Code E: USmin = 10 V, USmax = 32 V; "
                "Code F: USmin = 16 V, USmax = 32 V; etc.)",
            ):
                document.add_paragraph(line, style="List Bullet")
        document.add_paragraph(
            "The applicable supply voltage code for the DuT under test is specified in the "
            "individual test sections."
        )
        self._spacer(document)

    # ------------------------------------------------------------------
    # Section 5: Statement of Measurement Uncertainty
    # ------------------------------------------------------------------

    def _measurement_uncertainty(self, document):
        self._heading(document, "5. Statement of Measurement Uncertainty", 1)
        paragraph = document.add_paragraph()
        run = paragraph.add_run(
            "The data and results referenced in this document are true and accurate. The reader "
            "is cautioned that there may be errors within the calibration limits of the equipment "
            "and facilities. The measurement uncertainty was calculated for all measurements "
            "listed in this test report in accordance with GUM (‘Guide to the Expression of "
            "Uncertainty in Measurement’) and is documented in the quality system in "
            "accordance with DIN EN ISO/IEC 17025. Furthermore, component and process variability "
            "of devices similar to that tested may result in additional deviation. The "
            "manufacturer has the sole responsibility for the continued compliance of the device."
        )
        run.bold = True
        self._spacer(document)

    # ------------------------------------------------------------------
    # Section 6: Decision Rule & Conformity Assessment
    # ------------------------------------------------------------------

    def _decision_rule(self, document):
        self._heading(document, "6. Decision Rule & Conformity Assessment", 1)

        heading = document.add_paragraph()
        run = heading.add_run("Decision Rule")
        run.bold = True
        run.underline = True
        document.add_paragraph(
            "Measurement uncertainty is handled in accordance with GUM. All measurement device "
            "uncertainties used to evaluate conformity shall meet the requirements specified in "
            "Section 4 of this report."
        )

        heading = document.add_paragraph()
        run = heading.add_run("Conformity Assessment")
        run.bold = True
        run.underline = True
        document.add_paragraph(
            "Pass/fail determination is based on the functional status class (Class A to Class E, "
            "see Section 2) agreed between the test laboratory and the client for each function of "
            "the DuT. The tested sample may not represent the final production configuration. "
            "Agreed functions checked for the DuT are documented in Attachment II: Functional "
            "Checks for the DuT."
        )
        document.add_paragraph(
            "For each test in this report, the DuT was evaluated as “Pass – Met the "
            "criteria” or “Fail – Did not meet the criteria” in accordance with "
            "the requirements of ISO 16750-1:2023."
        )
        self._spacer(document)

    # ------------------------------------------------------------------
    # Section 7: Tests
    # ------------------------------------------------------------------

    def _tests_section(self, document, entries, category):
        self._heading(document, "7. Tests", 1)
        grouped = self._group_entries(entries, category)
        index = 0
        for group, group_entries in grouped.items():
            if category == "all":
                self._heading(document, group.title(), 2)
            for entry in group_entries:
                index += 1
                self._single_test(document, entry, index)

    def _single_test(self, document, entry, index):
        test = entry["test"]
        result = entry.get("result") or {}
        catalog = entry.get("catalog") or {}
        is_chemical = (test.get("category") or "").lower() == "chemical"
        measured = result.get("measured_values") or {}

        self._heading(document, f"7.{index} {test.get('test_name')}", 2)
        reference = document.add_paragraph()
        run = reference.add_run(f"Reference: {self._clause_reference(test)}")
        run.italic = True

        self._heading(document, "Test Method", 3)
        document.add_paragraph(
            catalog.get("purpose")
            or "Test method as defined by the applicable ISO 16750 clause and the project test specification."
        )

        schema = self._schema(catalog)
        rules_by_field = self._rules_by_field(catalog)
        if not schema:
            schema = [{"name": "measurement", "label": "Measurement"}]

        self._table_caption(document, f"Table 7.{index}.1: Test Parameters & Measured Values")
        parameters = document.add_table(rows=1, cols=3)
        self._style_table(parameters)
        self._header(parameters.rows[0], ["Parameter", "Requirements", "Measured / Applied Values"])
        for label, requirement, value in self._test_parameter_rows(test, catalog, schema, rules_by_field, measured, result):
            row = parameters.add_row().cells
            run = row[0].paragraphs[0].add_run(label)
            run.bold = True
            row[1].text = self._display(requirement)
            row[2].text = self._display(value)
        self._spacer(document)

        self._photo_for_test(document, entry, index)

        self._table_caption(document, "Laboratory Ambient Conditions")
        env = document.add_table(rows=2, cols=2)
        self._style_table(env)
        self._header(env.rows[0], ["Temperature:", "Humidity:"])
        env.rows[1].cells[0].text = self._display(result.get("temp") or "23 +/- 5 C")
        env.rows[1].cells[1].text = self._display(result.get("humidity") or "25 % - 75 %")
        self._spacer(document)

        if is_chemical:
            self._physical_check(document, measured)
            self._chemical_visual_check(document, result, measured)

        self._table_caption(document, "Acceptance Criteria / Result")
        acceptance = document.add_table(rows=2, cols=2)
        self._style_table(acceptance, header=False)
        acceptance.rows[0].cells[0].paragraphs[0].add_run("Acceptance Criteria").bold = True
        functional_status = test.get("functional_status") or "as agreed"
        status_prefix = functional_status if functional_status.lower().startswith("class") else f"Class {functional_status}"
        acceptance.rows[0].cells[1].text = self._display(
            f"Required Functional Status: {status_prefix}. "
            f"{test.get('acceptance_criteria') or catalog.get('purpose') or ''}".strip()
        )
        acceptance.rows[1].cells[0].paragraphs[0].add_run("Result").bold = True
        run = acceptance.rows[1].cells[1].paragraphs[0].add_run(self._result_statement(result))
        run.bold = True
        self._spacer(document)

        self._table_caption(document, "Evaluation Result")
        result_table = document.add_table(rows=0, cols=2)
        self._style_table(result_table)
        for label, value in (
            ("Evaluation Status", result.get("evaluation_status")),
            ("Evaluation Score", self._score(result.get("evaluation_score"))),
            ("Failed Rules", self._failed_rules(result)),
            ("Deviations", self._deviation(result)),
            ("Observations", result.get("observations")),
        ):
            self._field_row(result_table, label, value)
        self._spacer(document)

        if is_chemical:
            self._chemical_agent_table(document, measured)

        self._used_equipment(document, entry)

        technician = document.add_paragraph()
        technician.add_run("Test Technician: ").bold = True
        technician.add_run(self._display(result.get("engineer_name")))
        self._spacer(document)

    def _test_parameter_rows(self, test, catalog, schema, rules_by_field, measured, result):
        operating_mode_requirement = (catalog or {}).get("operating_mode") or test.get("operating_mode")
        rows = [
            ("Operating Mode", operating_mode_requirement, test.get("operating_mode")),
            ("Parameter Check Level", self._parameter_check_level(test), self._parameter_check_level(test)),
        ]
        for field in schema:
            name = field.get("name")
            label = field.get("label") or self._humanize_field_name(name)
            requirement = self._requirement_for_field(field, rules_by_field.get(name))
            rows.append((label, requirement, measured.get(name)))
        duration = test.get("duration_hours")
        rows.append(("Duration", f"{duration} h" if duration is not None else "-", f"{duration} h" if duration is not None else "-"))
        rows.append(("Sample Size", test.get("sample_size"), test.get("sample_size")))
        rows.append(("Sample No.", "-", "-"))
        rows.append(("Test Date(s)", "-", result.get("created_at") or date.today().isoformat()))
        return rows

    def _photo_for_test(self, document, entry, index):
        images = [a for a in entry.get("attachments", []) if is_image_attachment(a)]
        self._table_caption(document, f"Photo 7.{index}.1: Picture of the Test Setup")
        if not images:
            document.add_paragraph("No test setup photo has been uploaded for this test.")
            self._spacer(document)
            return
        for attachment in images:
            add_image_thumbnail_safe(document, attachment, width=2.5)
        self._spacer(document)

    def _physical_check(self, document, measured):
        visual_text = str(measured.get("visual_inspection") or "").lower()
        self._table_caption(document, "Physical Check")
        table = document.add_table(rows=1, cols=3)
        self._style_table(table)
        self._header(table.rows[0], ["Unit Material", "Check Criterion", "YES / NO"])

        row = table.add_row().cells
        row[0].paragraphs[0].add_run("Metallic Parts").bold = True
        row[1].text = (
            "A component shall be considered to pass this test if combine corrosion products "
            "(white and/or red) are not present over more than 30% of the combined primary "
            "surfaces after completion of the test."
        )
        row[2].text = "X" if "corrosion" in visual_text or "corroded" in visual_text else ""

        first = True
        for label, criterion, keywords in PHYSICAL_CHECK_ROWS:
            row = table.add_row().cells
            if first:
                row[0].paragraphs[0].add_run(label).bold = True
                first = False
            row[1].text = criterion
            row[2].text = "X" if visual_text and any(keyword in visual_text for keyword in keywords) else ""
        self._merge_first_column(table, start_row=2)
        self._spacer(document)

    @staticmethod
    def _merge_first_column(table, start_row):
        rows = table.rows[start_row:]
        if len(rows) < 2:
            return
        anchor = rows[0].cells[0]
        for row in rows[1:]:
            anchor = anchor.merge(row.cells[0])

    def _chemical_visual_check(self, document, result, measured):
        visual_text = str(measured.get("visual_inspection") or "").lower()
        status = result.get("evaluation_status")
        self._table_caption(document, "Visual Check / Parameter Check (Post-Test)")
        table = document.add_table(rows=1, cols=2)
        self._style_table(table)
        self._header(table.rows[0], ["Requirement", "Result"])
        corrosion_result = "Fail" if "corrosion" in visual_text else (status or "-")
        labelling_result = "Fail" if any(k in visual_text for k in ("label", "sticker", "paint")) else (status or "-")
        self._table_row(table, ("No detrimental corrosion.", corrosion_result))
        self._table_row(table, ("Marking and labelling shall remain visible and legible.", labelling_result))
        self._spacer(document)

    def _chemical_agent_table(self, document, measured):
        self._table_caption(document, "Chemical Agent Used")
        table = document.add_table(rows=1, cols=7)
        self._style_table(table)
        self._header(
            table.rows[0],
            ["ID No.", "Name of the Chemical Agent", "Brand", "Manufacturer", "First Date of Use", "Expire Date", "Concentration"],
        )
        agent = measured.get("chemical_agent")
        agent_id = "-"
        if agent:
            lowered = str(agent).lower()
            for keyword, identifier in CHEMICAL_AGENT_IDS:
                if keyword in lowered:
                    agent_id = identifier
                    break
        self._table_row(table, (agent_id, agent or "-", "-", "-", "-", "-", "-"))
        self._spacer(document)

    def _used_equipment(self, document, entry):
        self._table_caption(document, "Used Equipment")
        table = document.add_table(rows=1, cols=8)
        self._style_table(table)
        self._header(
            table.rows[0],
            ["Equipment No.", "Kind of Equipment", "Model Type", "Manufacturer", "Last Cal. Date", "Next Cal. Date", "Last Ver. Date", "Next Ver. Date"],
        )
        linked_equipment = entry.get("equipment") or []
        if linked_equipment:
            for item in linked_equipment:
                self._table_row(
                    table,
                    (
                        item.get("equipment_no"),
                        item.get("kind_of_equipment"),
                        item.get("model"),
                        item.get("manufacturer"),
                        item.get("last_calibration_date"),
                        item.get("next_calibration_date"),
                        item.get("last_verification_date"),
                        item.get("next_verification_date"),
                    ),
                )
        else:
            required_equipment = entry["test"].get("required_equipment") or []
            if isinstance(required_equipment, str):
                required_equipment = [required_equipment]
            for item in required_equipment:
                self._table_row(table, ("-", item, "-", "-", "-", "-", "-", "-"))
        self._spacer(document)

    # ------------------------------------------------------------------
    # Section 8: Comments
    # ------------------------------------------------------------------

    def _comments(self, document, entries, ai_sections):
        self._heading(document, "8. Comments", 1)
        comment_text = (ai_sections or {}).get("lab_comment")
        if not comment_text:
            statuses = [(entry.get("result") or {}).get("evaluation_status") for entry in entries]
            if statuses and all(status in ("PASS", None, "NOT EVALUATED") for status in statuses):
                comment_text = "The tested sample fulfilled the applicable requirements for the selected ISO 16750 test scope."
            else:
                comment_text = "One or more tests did not fulfill the applicable requirements. See individual test results."
        self._table_caption(document, "Comments")
        table = document.add_table(rows=0, cols=2)
        self._style_table(table)
        self._field_row(table, "Laboratory Comment", comment_text)
        self._field_row(table, "Reviewer Comment", "-")
        self._spacer(document)

    # ------------------------------------------------------------------
    # Attachments
    # ------------------------------------------------------------------

    def _attachments(self, document, entries, category):
        self._heading(document, "Attachment I: Photo Documentation", 1)
        self._photo_documentation(document, entries)
        self._spacer(document)
        self._heading(document, "Attachment II: Functional Checks for the DuT", 1)
        self._attachment_table(document, "Functional check records agreed between the test laboratory and the client.")
        self._spacer(document)
        self._heading(document, f"Attachment III: {self._attachment_iii_title(entries)}", 1)
        if self._has_chemical(entries):
            self._chemicals_versus_location(document, entries)
        else:
            self._attachment_table(document, "Tables and figures referenced from the applied ISO 16750 standards.")
        self._spacer(document)
        self._supporting_data_attachments(document, entries)

    def _supporting_data_attachments(self, document, entries):
        non_images = []
        for entry in entries:
            test_name = entry["test"].get("test_name") or "Test"
            for attachment in entry.get("attachments", []):
                if is_image_attachment(attachment):
                    continue
                item = dict(attachment)
                item["description"] = (
                    f"{test_name} - {item.get('description')}"
                    if item.get("description")
                    else test_name
                )
                non_images.append(item)
        if not non_images:
            self._attachment_table(document, "No supporting files have been uploaded.")
            return
        grouped = group_attachments_by_type(non_images)
        for attachment_type, attachments in grouped.items():
            title = ATTACHMENT_TYPE_LABELS.get(attachment_type, "Supporting Documents")
            self._table_caption(document, title)
            add_attachment_table(document, title, attachments)
            self._spacer(document)

    def _photo_documentation(self, document, entries):
        found = False
        for entry in entries:
            images = [
                attachment
                for attachment in entry.get("attachments", [])
                if is_image_attachment(attachment)
            ]
            if not images:
                continue
            found = True
            test = entry["test"]
            document.add_paragraph(test.get("test_name") or "Test").runs[0].bold = True
            for attachment in images:
                add_image_thumbnail_safe(document, attachment, width=1.6)
            self._spacer(document)
        if not found:
            self._attachment_table(document, "No test photo attachments have been uploaded.")

    def _chemicals_versus_location(self, document, entries):
        self._table_caption(document, "Chemicals versus Location")
        table = document.add_table(rows=1, cols=3)
        self._style_table(table)
        self._header(table.rows[0], ["Chemical Agent", "Mounting Location", "Test"])
        found = False
        for entry in entries:
            if (entry["test"].get("category") or "").lower() != "chemical":
                continue
            measured = (entry.get("result") or {}).get("measured_values") or {}
            agent = measured.get("chemical_agent")
            if not agent:
                continue
            found = True
            self._table_row(table, (agent, "-", entry["test"].get("test_name")))
        if not found:
            self._table_row(table, ("-", "-", "No chemical agent has been recorded."))
        self._spacer(document)

    # ------------------------------------------------------------------
    # Generic helpers
    # ------------------------------------------------------------------

    def _heading(self, document, text, level):
        if level > 0:
            self._spacer(document)
        paragraph = document.add_heading(level=level)
        run = paragraph.add_run(text)
        run.bold = True
        run.font.color.rgb = HEADER_BLUE

    @staticmethod
    def _spacer(document):
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(4)
        return paragraph

    @staticmethod
    def _page_break(document):
        document.add_page_break()

    @staticmethod
    def _table_caption(document, text):
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_before = Pt(4)
        paragraph.paragraph_format.space_after = Pt(2)
        run = paragraph.add_run(text)
        run.bold = True
        run.italic = True
        run.font.size = Pt(9)
        run.font.color.rgb = HEADER_BLUE

    def _style_table(self, table, header=True):
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        self._set_table_cell_margins(table)
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.space_after = Pt(0)
        if header and table.rows:
            for cell in table.rows[0].cells:
                self._set_cell_shading(cell, HEADER_FILL)

    @staticmethod
    def _set_table_cell_margins(table, top=80, bottom=80, left=120, right=120):
        tbl_pr = table._tbl.tblPr
        margins = OxmlElement("w:tblCellMar")
        for tag, value in (("top", top), ("bottom", bottom), ("left", left), ("right", right)):
            node = OxmlElement(f"w:{tag}")
            node.set(qn("w:w"), str(value))
            node.set(qn("w:type"), "dxa")
            margins.append(node)
        tbl_pr.append(margins)

    @staticmethod
    def _set_cell_shading(cell, hex_color):
        shading = OxmlElement("w:shd")
        shading.set(qn("w:val"), "clear")
        shading.set(qn("w:color"), "auto")
        shading.set(qn("w:fill"), hex_color)
        cell._tc.get_or_add_tcPr().append(shading)

    def _field_row(self, table, label, value):
        row = table.add_row().cells
        row[0].text = label
        row[1].text = self._display(value)
        row[0].paragraphs[0].runs[0].bold = True

    def _table_row(self, table, values):
        row = table.add_row().cells
        for cell, value in zip(row, values):
            cell.text = self._display(value)

    def _attachment_table(self, document, description):
        self._table_caption(document, "Attachment Placeholder")
        table = document.add_table(rows=1, cols=3)
        self._style_table(table)
        self._header(table.rows[0], ["Attachment", "Description", "Status"])
        self._table_row(table, ("Placeholder", description, "To be added"))
        self._spacer(document)

    def _header(self, row, values):
        for cell, value in zip(row.cells, values):
            cell.text = value
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True

    def _group_entries(self, entries, category):
        if category != "all":
            return OrderedDict([(category, entries)])
        grouped = OrderedDict()
        for entry in entries:
            key = (entry["test"].get("category") or "uncategorized").lower()
            grouped.setdefault(key, []).append(entry)
        return grouped

    @staticmethod
    def _add_page_field(paragraph, field_name):
        run = paragraph.add_run()
        fld_begin = OxmlElement("w:fldChar")
        fld_begin.set(qn("w:fldCharType"), "begin")
        instr_text = OxmlElement("w:instrText")
        instr_text.set(qn("xml:space"), "preserve")
        instr_text.text = field_name
        fld_end = OxmlElement("w:fldChar")
        fld_end.set(qn("w:fldCharType"), "end")
        run._r.append(fld_begin)
        run._r.append(instr_text)
        run._r.append(fld_end)

    @staticmethod
    def _schema(catalog):
        try:
            parsed = json.loads((catalog or {}).get("result_input_schema_json") or "[]")
        except (TypeError, ValueError):
            return []
        return parsed if isinstance(parsed, list) else []

    @staticmethod
    def _rules_by_field(catalog):
        try:
            parsed = json.loads((catalog or {}).get("evaluation_schema_json") or "{}")
        except (TypeError, ValueError):
            return {}
        rules = parsed.get("rules", []) if isinstance(parsed, dict) else []
        return {rule.get("field"): rule for rule in rules if isinstance(rule, dict) and rule.get("field")}

    def _requirement_for_field(self, field, rule):
        unit = f" {field.get('unit')}" if field.get("unit") else ""
        field_requirement = field.get("requirement") or field.get("required_value")
        if field_requirement:
            return str(field_requirement)
        if not rule:
            return "Mandatory recorded value" if field.get("required") else "Record if applicable"
        rule_type = rule.get("type")
        if rule_type == "range":
            return f"{self._display(rule.get('min'))}{unit} to {self._display(rule.get('max'))}{unit}"
        if rule_type == "min":
            return f">= {self._display(self._rule_value(rule, 'min'))}{unit}"
        if rule_type == "max":
            return f"<= {self._display(self._rule_value(rule, 'max'))}{unit}"
        if rule_type == "equals":
            return f"= {self._display(rule.get('value'))}{unit}"
        if rule_type == "not_equals":
            return f"Not {self._display(rule.get('value'))}{unit}"
        if rule_type == "contains":
            return f'Observation contains "{self._display(rule.get("value"))}"'
        if rule_type == "required":
            return str(field_requirement) if field_requirement else "Mandatory recorded value"
        return "Defined by catalog rule"

    @staticmethod
    def _failed_rules(result):
        details = result.get("evaluation_details") or {}
        failed = details.get("failed_rules") if isinstance(details, dict) else []
        if not failed:
            return "None"
        return "; ".join(rule.get("message", "") for rule in failed if rule.get("message"))

    @staticmethod
    def _category_title(category):
        labels = {
            "electrical": "Electrical Report",
            "environmental": "Environmental / Climatic Report",
            "chemical": "Chemical Report",
            "all": "Combined ISO 16750 Report",
        }
        return labels.get(category, "ISO 16750 Report")

    @staticmethod
    def _temperature_range(dut):
        if dut.get("temp_min") is None and dut.get("temp_max") is None:
            return "-"
        return f"{dut.get('temp_min')} C / {dut.get('temp_max')} C"

    @staticmethod
    def _test_dates(entries):
        dates = [
            (entry.get("result") or {}).get("created_at")
            for entry in entries
            if (entry.get("result") or {}).get("created_at")
        ]
        return ", ".join(dates) if dates else "-"

    @staticmethod
    def _test_standard(entries, category):
        if category != "all":
            return ISO_PART_FOR_CATEGORY.get(category, "ISO 16750:2023")
        categories = {
            (entry["test"].get("category") or "").lower()
            for entry in entries
            if entry["test"].get("category")
        }
        parts = sorted(
            ISO_PART_FOR_CATEGORY[key] for key in categories if key in ISO_PART_FOR_CATEGORY
        )
        return " / ".join(parts) if parts else "ISO 16750:2023"

    @staticmethod
    def _overall_result(entries):
        statuses = [(entry.get("result") or {}).get("evaluation_status") for entry in entries]
        if any(status == "FAIL" for status in statuses):
            return "FAIL"
        if any(status == "CONDITIONAL PASS" for status in statuses):
            return "CONDITIONAL PASS"
        if any(status == "PASS" for status in statuses):
            return "PASS"
        return "NOT EVALUATED"

    @staticmethod
    def _result_label(result):
        return result.get("evaluation_status") or result.get("result") or "ND"

    @staticmethod
    def _result_statement(result):
        status = (result.get("evaluation_status") or result.get("result") or "").upper()
        observations = result.get("observations")
        if status == "PASS":
            statement = "Pass – Met the criteria."
        elif status == "FAIL":
            statement = "Fail – Did not meet the criteria."
        else:
            statement = "Not Evaluated – Insufficient recorded data for automated assessment."
        if observations:
            statement = f"{statement} {observations}"
        return statement

    @staticmethod
    def _parameter_check_level(test):
        return "PC1"

    @staticmethod
    def _clause_reference(test):
        iso_part = TechnicalReportService._format_iso_part(test.get("iso_part"))
        clause_no = test.get("clause_no") or "-"
        return f"{iso_part}, Clause {clause_no}"

    @staticmethod
    def _format_iso_part(iso_part):
        if not iso_part:
            return "ISO 16750:2023"
        match = re.match(r"ISO\s*16750-?(\d+)", str(iso_part), re.IGNORECASE)
        if match:
            return f"ISO 16750-{match.group(1)}:2023"
        return str(iso_part)

    @staticmethod
    def _has_chemical(entries):
        return any((entry["test"].get("category") or "").lower() == "chemical" for entry in entries)

    @staticmethod
    def _attachment_iii_title(entries):
        if TechnicalReportService._has_chemical(entries):
            return "Chemicals versus Location"
        return "Tables & Figures from Applied Standards"

    @staticmethod
    def _score(score):
        return "-" if score is None else f"{score}%"

    @staticmethod
    def _deviation(result):
        if not result:
            return "-"
        if not result.get("has_deviation"):
            return "None"
        return result.get("deviation_description") or "Deviation recorded"

    @staticmethod
    def _display(value):
        if value is None or value == "":
            return "-"
        return str(value)

    @staticmethod
    def _rule_value(rule, key):
        return rule.get(key) if rule.get(key) is not None else rule.get("value")

    @staticmethod
    def _humanize_field_name(name):
        if not name:
            return "-"
        spaced = re.sub(r"(?<!^)(?=[A-Z])", " ", str(name))
        return spaced.replace("_", " ").title()
