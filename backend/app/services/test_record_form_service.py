"""DOCX generator for structured ISO 16750 test record forms."""

import io
import json
from datetime import date

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from ..utils.attachment_docx import (
    ATTACHMENT_TYPE_LABELS,
    add_attachment_table,
    add_image_thumbnail_safe,
    group_attachments_by_type,
    is_image_attachment,
)


DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
HEADER_BLUE = RGBColor(0x1F, 0x4E, 0x79)


class TestRecordFormService:
    """Builds structured test record form DOCX files."""

    def build_docx(
        self,
        dut,
        test,
        plan_item=None,
        catalog=None,
        result=None,
        attachments=None,
        equipment=None,
    ):
        """Return a BytesIO DOCX buffer for one test record form."""
        document = Document()
        self._configure_document(document)

        self._add_title_block(document, test)
        self._add_front_matter(document, dut, test, plan_item, result)
        self._add_equipment_table(document, test, equipment or [])
        self._add_environmental_conditions(document, result)

        document.add_page_break()
        self._add_test_parameters(document, test, plan_item, catalog, result)
        self._add_acceptance_table(document, test, result)
        self._add_requirements_results(document, test, plan_item, result)
        self._add_technician_comments(document, result)

        document.add_page_break()
        self._add_supervisor_page(document, result, attachments or [])

        buffer = io.BytesIO()
        document.save(buffer)
        buffer.seek(0)
        return buffer

    def _configure_document(self, document):
        section = document.sections[0]
        section.top_margin = Inches(0.55)
        section.bottom_margin = Inches(0.55)
        section.left_margin = Inches(0.55)
        section.right_margin = Inches(0.55)

        style = document.styles["Normal"]
        style.font.name = "Arial"
        style.font.size = Pt(9)
        self._header_footer(section)

    def _header_footer(self, section):
        header = section.header.paragraphs[0]
        header.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = header.add_run("FR 14 03 - ISO 16750 TEST RECORD FORM")
        run.bold = True
        run.font.size = Pt(9)

        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer.add_run("Controlled laboratory form - Page ")
        self._add_page_number(footer)

    def _add_title_block(self, document, test):
        table = document.add_table(rows=2, cols=3)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        self._style_table(table)
        widths = (1.5, 4.5, 1.6)
        for row in table.rows:
            for idx, cell in enumerate(row.cells):
                cell.width = Inches(widths[idx])
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

        table.cell(0, 0).text = "FR 14 03"
        title_cell = table.cell(0, 1)
        title_cell.text = "TEST RECORD FORM"
        title_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_cell.paragraphs[0].runs[0].bold = True
        title_cell.paragraphs[0].runs[0].font.size = Pt(14)
        table.cell(0, 2).text = "Page 1 / 3"
        table.cell(1, 0).text = "Standard"
        table.cell(1, 1).text = test.get("standard_reference") or "ISO 16750"
        table.cell(1, 2).text = date.today().isoformat()
        document.add_paragraph()

        title = document.add_table(rows=1, cols=1)
        self._style_table(title)
        title.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell = title.cell(0, 0)
        cell.text = test.get("test_name") or "ISO 16750 Test"
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.runs[0].bold = True
        paragraph.runs[0].font.size = Pt(16)
        document.add_paragraph()

    def _add_front_matter(self, document, dut, test, plan_item, result):
        self._table_caption(document, "Table 1: DUT Identity and Test Requirement Information")
        rows = [
            ("Manufacturer", dut.get("manufacturer")),
            ("Project No.", dut.get("project")),
            ("Product Description", dut.get("name")),
            ("Model Name", dut.get("part_number")),
            ("Supply Voltage", dut.get("nominal_voltage")),
            ("Vehicle Type Classification", "-"),
            ("Component Classification by Location", dut.get("mounting_location")),
            ("Electric Operating Conditions", self._first(test, plan_item, "operating_mode")),
            ("DUT Test Requirement Level", self._first(test, plan_item, "required_test_level")),
            ("Function Type", "-"),
            ("Test Sample Classification", self._first(test, plan_item, "sample_size")),
            ("Deviation/s", self._deviation_text(result)),
            ("Sample No.", self._first(test, plan_item, "sample_size")),
            ("Test Location", "-"),
            ("Shelf", "-"),
        ]
        table = document.add_table(rows=0, cols=2)
        self._style_table(table)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for label, value in rows:
            row = table.add_row().cells
            row[0].text = f"{label.upper()}:"
            row[1].text = self._display(value)
            row[0].paragraphs[0].runs[0].bold = True
        self._spacer(document)

    def _add_equipment_table(self, document, test, linked_equipment):
        self._spacer(document)
        self._section_heading(
            document,
            "Equipment used: (first, check if the calibration/verification dates are valid)",
        )
        self._table_caption(document, "Table 2: Equipment Used")
        table = document.add_table(rows=9, cols=10)
        self._style_table(table)
        headers = [
            "Equipment No",
            "Kind of equipment",
            "Model",
            "Type",
            "Manufacturer",
            "Last Cal Date",
            "Next Cal Date",
            "Last Ver Date",
            "Next Ver Date",
            "Using Status",
        ]
        self._set_header_row(table.rows[0], headers)
        if linked_equipment:
            for index, item in enumerate(linked_equipment[:8], start=1):
                row = table.rows[index].cells
                values = [
                    item.get("equipment_no"),
                    item.get("kind_of_equipment"),
                    item.get("model"),
                    item.get("type"),
                    item.get("manufacturer"),
                    item.get("last_calibration_date"),
                    item.get("next_calibration_date"),
                    item.get("last_verification_date"),
                    item.get("next_verification_date"),
                    item.get("using_status"),
                ]
                for cell, value in zip(row, values):
                    cell.text = self._display(value)
            self._spacer(document)
            return
        equipment = test.get("required_equipment") or []
        if isinstance(equipment, str):
            equipment = [equipment]
        for index, item in enumerate(equipment[:8], start=1):
            row = table.rows[index].cells
            row[0].text = str(index)
            row[1].text = str(item)
        self._spacer(document)

    def _add_environmental_conditions(self, document, result):
        self._spacer(document)
        self._section_heading(document, "Environmental Conditions")
        self._table_caption(document, "Table 3: Environmental Conditions")
        table = document.add_table(rows=2, cols=3)
        self._style_table(table)
        self._set_header_row(table.rows[0], ["Temperature", "Humidity", "Air Pressure"])
        row = table.rows[1].cells
        row[0].text = self._with_unit(result.get("temp") if result else None, "C")
        row[1].text = self._with_unit(result.get("humidity") if result else None, "%")
        row[2].text = "-"
        self._spacer(document)

    def _add_test_parameters(self, document, test, plan_item, catalog, result):
        self._section_heading(document, "Table: Test Parameters")
        self._table_caption(document, "Table 4: Test Parameters")
        schema = self._result_schema(catalog)
        rules_by_field = self._rules_by_field(catalog)
        measured = (result or {}).get("measured_values") or {}
        table = document.add_table(rows=1, cols=3)
        self._style_table(table)
        self._set_header_row(table.rows[0], ["", "Requirements", "Measured/Applied Values"])
        fixed_rows = [
            ("Component Operating Condition", self._first(test, plan_item, "operating_mode")),
            ("DUT Test Required Level", self._first(test, plan_item, "required_test_level")),
            ("Required Functional Status", self._first(test, plan_item, "functional_status")),
            ("Sample No.", self._first(test, plan_item, "sample_size")),
            ("Test Date(s)", date.today().isoformat()),
        ]
        for label, requirement in fixed_rows:
            row = table.add_row().cells
            row[0].text = label
            row[1].text = self._display(requirement)
            row[2].text = self._display(requirement)
        if not schema:
            schema = [{"name": "measurement", "label": "Measurement", "unit": ""}]
        for field in schema:
            row = table.add_row().cells
            label = field.get("label") or field.get("name", "")
            row[0].text = label
            row[1].text = self._requirement_for_field(field, rules_by_field.get(field.get("name")))
            row[2].text = self._display(measured.get(field.get("name")))

        evaluation = (result or {}).get("evaluation_details") or {}
        for label, value in (
            ("Evaluation Status", (result or {}).get("evaluation_status")),
            ("Evaluation Score", self._score_text((result or {}).get("evaluation_score"))),
            ("Failed Rules", self._failed_rules_text(evaluation)),
        ):
            row = table.add_row().cells
            row[0].text = label
            row[1].text = label
            row[2].text = self._display(value)
        self._spacer(document)

    def _add_acceptance_table(self, document, test, result):
        self._spacer(document)
        self._table_caption(document, "Table 5: Acceptance Criteria")
        table = document.add_table(rows=2, cols=2)
        self._style_table(table)
        table.rows[0].cells[0].text = "Acceptance Criteria"
        table.rows[0].cells[1].text = self._display(test.get("acceptance_criteria"))
        table.rows[1].cells[0].text = "Result"
        table.rows[1].cells[1].text = self._display(
            (result or {}).get("evaluation_status") or (result or {}).get("result")
        )
        self._spacer(document)

    def _add_requirements_results(self, document, test, plan_item, result):
        self._spacer(document)
        self._section_heading(document, "Requirements & Results")
        self._table_caption(document, "Table 6: Requirements and Results")
        table = document.add_table(rows=2, cols=7)
        self._style_table(table)
        headers1 = [
            "Test No.",
            "ISO 16750 Clause No.",
            "Name of the Test",
            "Required Functional Status",
            "Sample No.",
            "Deviations",
            "Result",
        ]
        headers2 = [
            "Test No.",
            "ISO 16750 Clause No.",
            "Operating Condition / Required Test Level",
            "Required Functional Status",
            "Sample No.",
            "Deviations",
            "Result",
        ]
        self._set_header_row(table.rows[0], headers1)
        self._set_header_row(table.rows[1], headers2)
        row = table.add_row().cells
        row[0].text = self._display((plan_item or {}).get("planned_test_no"))
        row[1].text = self._display(self._first(test, plan_item, "clause_no"))
        row[2].text = self._display(self._first(test, plan_item, "test_name"))
        row[3].text = self._display(self._first(test, plan_item, "functional_status"))
        row[4].text = self._display(self._first(test, plan_item, "sample_size"))
        row[5].text = self._deviation_text(result)
        row[6].text = self._display((result or {}).get("evaluation_status") or (result or {}).get("result"))

        self._spacer(document)
        self._table_caption(document, "Table 7: Operating Condition and Required Test Level")
        table2 = document.add_table(rows=1, cols=4)
        self._style_table(table2)
        self._set_header_row(
            table2.rows[0],
            ["ISO Part", "Operating Condition", "Required Test Level", "Test Date(s)"],
        )
        row2 = table2.add_row().cells
        row2[0].text = self._display(self._first(test, plan_item, "iso_part"))
        row2[1].text = self._display(self._first(test, plan_item, "operating_mode"))
        row2[2].text = self._display(self._first(test, plan_item, "required_test_level"))
        row2[3].text = date.today().isoformat()
        self._spacer(document)

    def _add_technician_comments(self, document, result):
        self._spacer(document)
        self._section_heading(document, "Comments by Test Technician")
        self._table_caption(document, "Table 8: Test Technician Comments")
        table = document.add_table(rows=2, cols=1)
        self._style_table(table)
        table.rows[0].cells[0].text = "Technician Comments"
        table.rows[0].cells[0].paragraphs[0].runs[0].bold = True
        table.rows[1].cells[0].text = self._display((result or {}).get("observations"))
        self._spacer(document)

    def _add_supervisor_page(self, document, result, attachments):
        self._section_heading(document, "Comments by Supervisor")
        self._table_caption(document, "Table 9: Supervisor Comments")
        supervisor = document.add_table(rows=2, cols=1)
        self._style_table(supervisor)
        supervisor.rows[0].cells[0].text = "Supervisor Comments"
        supervisor.rows[0].cells[0].paragraphs[0].runs[0].bold = True
        supervisor.rows[1].cells[0].text = ""
        self._spacer(document)
        self._section_heading(document, "Important Notes")
        self._table_caption(document, "Table 10: Important Notes")
        notes = [
            "All measurement equipment calibration status shall be verified before testing.",
            "Any deviation from the approved method shall be documented in this form.",
            "The final result shall be reviewed together with raw measurement records.",
        ]
        table = document.add_table(rows=1, cols=2)
        self._style_table(table)
        self._set_header_row(table.rows[0], ["No.", "Note"])
        for index, note in enumerate(notes, start=1):
            row = table.add_row().cells
            row[0].text = str(index)
            row[1].text = note
        self._spacer(document)

        self._add_attachment_references(document, attachments)

        self._spacer(document)
        self._section_heading(document, "Signature Area")
        self._table_caption(document, "Table 11: Signatures")
        signatures = document.add_table(rows=4, cols=2)
        self._style_table(signatures)
        self._set_header_row(signatures.rows[0], ["Test Technician", "Supervisor"])
        signatures.rows[1].cells[0].text = self._display((result or {}).get("engineer_name"))
        signatures.rows[1].cells[1].text = ""
        signatures.rows[2].cells[0].text = "Signature:"
        signatures.rows[2].cells[1].text = "Signature:"
        signatures.rows[3].cells[0].text = "Date:"
        signatures.rows[3].cells[1].text = "Date:"
        self._spacer(document)

    def _add_attachment_references(self, document, attachments):
        self._section_heading(document, "Attachment References")
        grouped = group_attachments_by_type(attachments)
        ordered_types = [
            "test_photo",
            "measurement_file",
            "temperature_humidity_log",
            "calibration_certificate",
            "raw_data",
            "supporting_document",
        ]
        if not attachments:
            document.add_paragraph("No attachments have been uploaded for this test.")
            self._spacer(document)
            return
        for attachment_type in ordered_types:
            items = grouped.get(attachment_type, [])
            if not items:
                continue
            title = ATTACHMENT_TYPE_LABELS.get(attachment_type, "Supporting Documents")
            images = [item for item in items if is_image_attachment(item)]
            non_images = [item for item in items if not is_image_attachment(item)]
            if images:
                document.add_paragraph(title).runs[0].bold = True
                for attachment in images:
                    add_image_thumbnail_safe(document, attachment, width=1.35)
                self._spacer(document)
            if non_images:
                add_attachment_table(document, title, non_images)
                self._spacer(document)

    def _section_heading(self, document, text):
        paragraph = document.add_paragraph()
        run = paragraph.add_run(text)
        run.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = HEADER_BLUE

    @staticmethod
    def _spacer(document):
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(4)
        return paragraph

    @staticmethod
    def _table_caption(document, text):
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_before = Pt(4)
        paragraph.paragraph_format.space_after = Pt(2)
        run = paragraph.add_run(text)
        run.bold = True
        run.italic = True
        run.font.size = Pt(9)
        run.font.color.rgb = HEADER_BLUE

    @staticmethod
    def _style_table(table):
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.space_after = Pt(0)

    def _set_header_row(self, row, headers):
        for cell, header in zip(row.cells, headers):
            cell.text = header
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True

    @staticmethod
    def _add_page_number(paragraph):
        run = paragraph.add_run()
        fld_begin = OxmlElement("w:fldChar")
        fld_begin.set(qn("w:fldCharType"), "begin")
        instr_text = OxmlElement("w:instrText")
        instr_text.set(qn("xml:space"), "preserve")
        instr_text.text = "PAGE"
        fld_end = OxmlElement("w:fldChar")
        fld_end.set(qn("w:fldCharType"), "end")
        run._r.append(fld_begin)
        run._r.append(instr_text)
        run._r.append(fld_end)

    def _result_schema(self, catalog):
        if not catalog:
            return []
        try:
            schema = json.loads(catalog.get("result_input_schema_json") or "[]")
        except (TypeError, ValueError):
            return []
        return schema if isinstance(schema, list) else []

    def _rules_by_field(self, catalog):
        if not catalog:
            return {}
        try:
            parsed = json.loads(catalog.get("evaluation_schema_json") or "{}")
        except (TypeError, ValueError):
            return {}
        rules = parsed.get("rules", []) if isinstance(parsed, dict) else []
        return {rule.get("field"): rule for rule in rules if isinstance(rule, dict) and rule.get("field")}

    def _requirement_for_field(self, field, rule):
        unit = f" {field.get('unit')}" if field.get("unit") else ""
        field_requirement = field.get("requirement") or field.get("required_value")
        if not rule:
            if field_requirement:
                return str(field_requirement)
            return "Mandatory recorded value" if field.get("required") else "Record if applicable"
        rule_type = rule.get("type")
        if rule_type == "range":
            return f"{self._display(rule.get('min'))}{unit} to {self._display(rule.get('max'))}{unit}"
        if rule_type == "min":
            return f">= {self._display(self._rule_value(rule, 'min'))}{unit}"
        if rule_type == "max":
            return f"<= {self._display(self._rule_value(rule, 'max'))}{unit}"
        if rule_type == "equals":
            return f"= {self._display(rule.get('value'))}{unit}"
        if rule_type == "not_equals":
            return f"Not {self._display(rule.get('value'))}{unit}"
        if rule_type == "contains":
            return f'Observation contains "{self._display(rule.get("value"))}"'
        if rule_type == "required":
            return str(field_requirement) if field_requirement else "Mandatory recorded value"
        return "Defined by catalog rule"

    def _deviation_text(self, result):
        if not result:
            return "-"
        if not result.get("has_deviation"):
            return "None"
        return result.get("deviation_description") or "Deviation recorded"

    def _failed_rules_text(self, evaluation):
        failed = evaluation.get("failed_rules") if isinstance(evaluation, dict) else []
        if not failed:
            return "None"
        return "; ".join(rule.get("message", "") for rule in failed if rule.get("message"))

    @staticmethod
    def _score_text(score):
        if score is None:
            return "-"
        return f"{score}%"

    @staticmethod
    def _with_unit(value, unit):
        if value is None or value == "":
            return "-"
        return f"{value} {unit}"

    @staticmethod
    def _display(value):
        if value is None or value == "":
            return "-"
        return str(value)

    @staticmethod
    def _rule_value(rule, key):
        return rule.get(key) if rule.get(key) is not None else rule.get("value")

    @staticmethod
    def _first(*sources_and_key):
        *sources, key = sources_and_key
        for source in sources:
            if source and source.get(key):
                return source.get(key)
        return None
