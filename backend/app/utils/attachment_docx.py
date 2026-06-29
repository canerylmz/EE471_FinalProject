"""Safe DOCX helpers for listing and embedding uploaded attachments."""

from collections import defaultdict
from pathlib import Path

from docx.shared import Inches


ATTACHMENT_TYPE_LABELS = {
    "test_photo": "Test Photos",
    "measurement_file": "Measurement Files",
    "temperature_humidity_log": "Temperature-Humidity Logs",
    "calibration_certificate": "Calibration Certificates",
    "raw_data": "Raw Data",
    "test_record_form": "Test Record Forms",
    "supporting_document": "Supporting Documents",
}
IMAGE_EXTENSIONS = {"png", "jpg", "jpeg"}


def group_attachments_by_type(attachments):
    """Return attachments grouped by attachment_type."""
    grouped = defaultdict(list)
    for attachment in attachments or []:
        grouped[attachment.get("attachment_type") or "supporting_document"].append(attachment)
    return dict(grouped)


def is_image_attachment(attachment):
    """Return True when an attachment points to a supported image extension."""
    filename = attachment.get("original_filename") or attachment.get("file_name") or ""
    extension = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
    return extension in IMAGE_EXTENSIONS


def add_attachment_table(document, title, attachments):
    """Add a simple attachment reference table to a DOCX document."""
    document.add_paragraph(title).runs[0].bold = True
    table = document.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    headers = ["Filename", "Type", "Description", "Uploaded At"]
    for cell, header in zip(table.rows[0].cells, headers):
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True
    if not attachments:
        row = table.add_row().cells
        row[0].text = "No attachments"
        row[1].text = "-"
        row[2].text = "-"
        row[3].text = "-"
        return table

    for attachment in attachments:
        row = table.add_row().cells
        row[0].text = attachment.get("original_filename") or attachment.get("file_name") or "-"
        row[1].text = ATTACHMENT_TYPE_LABELS.get(
            attachment.get("attachment_type"), "Supporting Documents"
        )
        row[2].text = attachment.get("description") or "-"
        row[3].text = attachment.get("created_at") or "-"
    return table


def add_image_thumbnail_safe(document, attachment, width=1.6):
    """Embed a thumbnail if possible; otherwise list the image file without raising."""
    filename = attachment.get("original_filename") or attachment.get("file_name") or "image"
    description = attachment.get("description") or ""
    file_path = Path(attachment.get("file_path") or "")
    paragraph = document.add_paragraph()
    paragraph.add_run(filename).bold = True
    if description:
        paragraph.add_run(f" - {description}")
    if not file_path.exists() or not is_image_attachment(attachment):
        document.add_paragraph("Image file is not available for embedding.")
        return False
    try:
        document.add_picture(str(file_path), width=Inches(width))
        return True
    except Exception:
        document.add_paragraph("Image could not be embedded; file is referenced by filename.")
        return False
