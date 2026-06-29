"""python-docx based exporters for the test plan and formal report."""

import io
from datetime import date

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from .attachment_docx import add_attachment_table, add_image_thumbnail_safe, is_image_attachment
from .qr_helper import generate_qr_image

NAVY = RGBColor(0x0F, 0x17, 0x2A)
AMBER = RGBColor(0xF5, 0x9E, 0x0B)

REPORT_SECTION_TITLES = {
    "test_amaci": "1. Test Purpose",
    "test_kosullari": "2. Test Conditions",
    "olcum_sonuclari": "3. Measurement Results",
    "gozlemler": "4. Observations",
    "kabul_degerlendirme": "5. Acceptance Criteria Evaluation",
    "sonuc": "6. Result and Decision",
    "sapma_analizi": "7. Deviation Analysis and Corrective Action",
}


def _add_page_number(paragraph):
    """Insert a Word PAGE field into `paragraph`."""
    run = paragraph.add_run()

    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")

    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"

    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")

    run._r.append(fld_begin)
    run._r.append(instr_text)
    run._r.append(fld_end)


def _add_toc(document):
    """Insert a Word TOC field that can be refreshed in Word."""
    paragraph = document.add_paragraph()
    run = paragraph.add_run()

    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")

    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = 'TOC \\o "1-2" \\h \\z \\u'

    fld_separate = OxmlElement("w:fldChar")
    fld_separate.set(qn("w:fldCharType"), "separate")

    placeholder = OxmlElement("w:t")
    placeholder.text = "Right-click this field in Word and select 'Update Field' to refresh the contents."

    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")

    run._r.append(fld_begin)
    run._r.append(instr_text)
    run._r.append(fld_separate)
    run._r.append(placeholder)
    run._r.append(fld_end)


def _add_footer(document):
    section = document.sections[0]
    paragraph = section.footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run("TestForge - Page ")
    _add_page_number(paragraph)


def _add_field_row(table, label, value):
    row = table.add_row().cells
    row[0].text = label
    run = row[0].paragraphs[0].runs[0]
    run.bold = True
    row[1].text = "" if value is None else str(value)


def build_plan_docx(dut, tests):
    """Build the ISO 16750 test plan .docx and return a BytesIO buffer."""
    document = Document()
    _add_footer(document)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("TestForge")
    run.bold = True
    run.font.size = Pt(40)
    run.font.color.rgb = AMBER

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("ISO 16750 Automotive Electronics Test Plan")
    run.font.size = Pt(16)
    run.font.color.rgb = NAVY

    dut_title = document.add_paragraph()
    dut_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = dut_title.add_run(dut["name"])
    run.bold = True
    run.font.size = Pt(22)

    document.add_paragraph()

    info_table = document.add_table(rows=0, cols=2)
    info_table.style = "Light Grid Accent 1"
    for label, value in (
        ("Manufacturer", dut.get("manufacturer", "")),
        ("Part Number", dut.get("part_number", "")),
        ("Client", dut.get("customer", "")),
        ("Project", dut.get("project", "")),
        ("Mounting Location", dut.get("mounting_location", "")),
        ("Nominal Voltage", dut.get("nominal_voltage", "")),
        ("Power Class", dut.get("power_class", "")),
        ("Operating Temperature", f"{dut.get('temp_min')} C / {dut.get('temp_max')} C"),
        ("IP Code", dut.get("ip_class", "")),
        ("Report Date", date.today().isoformat()),
    ):
        _add_field_row(info_table, label, value)

    document.add_paragraph()
    qr_buffer = generate_qr_image(f"TestForge|DUT:{dut['id']}|{dut['name']}")
    qr_paragraph = document.add_paragraph()
    qr_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    qr_run = qr_paragraph.add_run()
    qr_run.add_picture(qr_buffer, width=Inches(1.2))

    document.add_page_break()

    document.add_heading("Contents", level=1)
    _add_toc(document)
    document.add_page_break()

    for index, test in enumerate(tests, start=1):
        document.add_heading(f"{index}. {test['test_name']}", level=1)

        details = document.add_table(rows=0, cols=2)
        details.style = "Light List Accent 1"
        equipment = test.get("required_equipment") or []
        if not isinstance(equipment, list):
            equipment = [equipment]

        for label, value in (
            ("Standard Reference", test.get("standard_reference", "")),
            ("Category", test.get("category", "")),
            ("Status", test.get("status", "")),
            ("Duration (hours)", test.get("duration_hours", "")),
            ("Severity Level", test.get("severity_level", "")),
        ):
            _add_field_row(details, label, value)

        document.add_paragraph()
        document.add_heading("Required Equipment", level=2)
        if equipment:
            for item in equipment:
                document.add_paragraph(str(item), style="List Bullet")
        else:
            document.add_paragraph("Not specified")

        document.add_heading("Acceptance Criteria", level=2)
        document.add_paragraph(test.get("acceptance_criteria", "Not specified"))

        if index < len(tests):
            document.add_page_break()

    buffer = io.BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer


def build_report_docx(dut, test, result, report, attachments=None):
    """Build the formal test report .docx and return a BytesIO buffer."""
    document = Document()
    _add_footer(document)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("TestForge")
    run.bold = True
    run.font.size = Pt(32)
    run.font.color.rgb = AMBER

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Laboratory Test Report")
    run.font.size = Pt(16)
    run.font.color.rgb = NAVY

    document.add_paragraph()

    info_table = document.add_table(rows=0, cols=2)
    info_table.style = "Light Grid Accent 1"
    for label, value in (
        ("DUT", dut.get("name", "")),
        ("Part Number", dut.get("part_number", "")),
        ("Test", test.get("test_name", "")),
        ("Standard Reference", test.get("standard_reference", "")),
        ("Result", result.get("result", "")),
        ("Engineer", result.get("engineer_name", "")),
        ("Date", date.today().isoformat()),
    ):
        _add_field_row(info_table, label, value)

    document.add_paragraph()

    measured_values = result.get("measured_values") or {}
    if measured_values:
        document.add_heading("Measured Values", level=2)
        mv_table = document.add_table(rows=1, cols=2)
        mv_table.style = "Light List Accent 1"
        header_cells = mv_table.rows[0].cells
        header_cells[0].text = "Parameter"
        header_cells[1].text = "Value"
        for key, value in measured_values.items():
            row = mv_table.add_row().cells
            row[0].text = str(key)
            row[1].text = str(value)
        document.add_paragraph()

    sections = report.get("sections", {})
    for key, heading in REPORT_SECTION_TITLES.items():
        if key in sections:
            document.add_heading(heading, level=1)
            document.add_paragraph(sections[key])

    attachments = attachments or []
    if attachments:
        document.add_heading("Attachments and Photo Evidence", level=1)
        image_attachments = [item for item in attachments if is_image_attachment(item)]
        other_attachments = [item for item in attachments if not is_image_attachment(item)]
        if image_attachments:
            document.add_heading("Test Photos", level=2)
            for attachment in image_attachments:
                add_image_thumbnail_safe(document, attachment, width=2.0)
        if other_attachments:
            add_attachment_table(document, "Supporting Files", other_attachments)

    buffer = io.BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer
